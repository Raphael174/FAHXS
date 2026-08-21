"""Version française du document de cadre théorique du solveur stationnaire
calandre-et-tubes (.docx).

Document de référence pour les résultats Hélium / LN2 / Eau. À régénérer
après toute modification des équations de conservation ou des corrélations de
fermeture :

    python docs/generate_shelltube_theory_docx_fr.py

Toute la machinerie (classe ``Doc``, convertisseur LaTeX -> OMML, numérotation
des équations, renvois en deux passes) est importée de la version anglaise
``generate_shelltube_theory_docx.py`` : seul le texte est dupliqué. Les
équations, leurs étiquettes et leur numérotation sont identiques dans les deux
versions, de sorte qu'un renvoi « Eq. 17 » désigne la même relation dans l'une
et l'autre.
"""
from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from generate_shelltube_theory_docx import (  # noqa: E402
    COLD,
    FLAG,
    HEAD_FONT,
    INK,
    SOFT,
    Doc,
)
from docx.shared import Pt  # noqa: E402


def build(fig_dir: str, refmap=None) -> Doc:
    D = Doc(refmap)

    # ---------------- page de titre ----------------
    t = D.d.add_paragraph()
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("Solveur stationnaire calandre-et-tubes")
    r.font.name = HEAD_FONT; r.font.size = Pt(22); r.bold = True; r.font.color.rgb = INK

    s = D.d.add_paragraph()
    s.paragraph_format.space_after = Pt(10)
    r = s.add_run("Cadre théorique, équations de conservation et corrélations empiriques")
    r.font.name = HEAD_FONT; r.font.size = Pt(13); r.font.color.rgb = SOFT

    s = D.d.add_paragraph()
    s.paragraph_format.space_after = Pt(14)
    r = s.add_run("Document de référence pour les résultats Hélium / LN₂ / Eau")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = COLD

    D.table(
        ["Élément", "Valeur"],
        [["Modèle", "Échangeur de chaleur de combusteur calandre-et-tubes, quasi-1D, stationnaire"],
         ["Configuration", "Gaz de combustion à l'intérieur des tubes droits ; réfrigérant en écoulement transversal côté calandre"],
         ["Implémentation principale", "1Dmodel/main_solve_shellntube.py"],
         ["Modèle de paroi", "physics/heat_conduction.py — OneDimensionalSteadyConduction_ShellnHelicalTube"],
         ["Fermeture côté calandre", "physics/bell_delaware.py (monophasique) ; physics/liquid_flow/ (diphasique, supercritique)"],
         ["Point d'entrée", "1Dmodel/main_steady.py — combustorProp.HX_config = \"shellntube\""],
         ["Unités", "SI partout, à l'exception des tables de propriétés matériaux qui sont en °C"]],
        "Périmètre du document et source de référence. Régénérer ce document avec "
        "`docs/generate_shelltube_theory_docx_fr.py` après toute modification des équations ou des fermetures.",
        widths=[4.2, 12.0])

    # ---------------- 1 ----------------
    D.h1("1.  Périmètre et configuration physique")

    D.p("Les produits de combustion issus d'un combusteur ergol liquide/oxygène circulent axialement "
        "à l'intérieur d'un faisceau de tubes droits. Le réfrigérant s'écoule côté calandre, contraint "
        "à un trajet transversal en zigzag par des chicanes segmentées. La chaleur traverse "
        "radialement la paroi du tube, du gaz vers le réfrigérant.")

    D.p("Ce document spécifie le modèle tel qu'implémenté, y compris les corrélations empiriques et "
        "les hypothèses attachées à chacune. Il constitue la base technique commune aux résultats "
        "Hélium, LN₂/N₂ supercritique et Eau, qui ne diffèrent que par la branche de fermeture "
        "côté réfrigérant sollicitée : la structure côté gaz, la paroi et l'itération sont "
        "identiques pour les trois. Les conséquences propres à chaque fluide sont rassemblées en "
        "section 10.")

    D.h2("1.1  Nomenclature")
    D.table(
        ["Symbole", "Grandeur", "Unités"],
        [["A_i, A_hot, A_cold", "Section de passage interne ; aires nodales des faces chaude et froide", "m²"],
         ["B", "Espacement des chicanes", "m"],
         ["Bo", "Nombre d'ébullition, q″/(G h_fg)", "–"],
         ["cp", "Chaleur massique à pression constante", "J/(kg·K)"],
         ["D_i, D_o", "Diamètres intérieur et extérieur du tube", "m"],
         ["D_s, D_otl", "Diamètre intérieur de calandre ; diamètre d'enveloppe du faisceau", "m"],
         ["dQ_i", "Puissance thermique au nœud i, par tube", "W"],
         ["f", "Coefficient de frottement de Darcy", "–"],
         ["G_s, G_w", "Flux massiques transversal et en fenêtre, côté calandre", "kg/(m²·s)"],
         ["h", "Enthalpie massique", "J/kg"],
         ["h_g, h_c", "Coefficients d'échange côté gaz et côté réfrigérant", "W/(m²·K)"],
         ["h_fg", "Chaleur latente de vaporisation", "J/kg"],
         ["j", "Facteur j de Colburn", "–"],
         ["k_w", "Conductivité thermique de la paroi", "W/(m·K)"],
         ["L_tube", "Longueur de tube", "m"],
         ["ṁ_g, ṁ_c, ṁ_tube", "Débits massiques gaz chaud, réfrigérant et par tube", "kg/s"],
         ["N, N_tubes, N_b", "Nœuds axiaux ; tubes du faisceau ; chicanes", "–"],
         ["N_tcc, N_tcw", "Rangées traversées par section transversale ; par fenêtre", "–"],
         ["Nu, Pr, Re", "Nombres de Nusselt, Prandtl, Reynolds", "–"],
         ["p, p_r", "Pression ; pression réduite p/p_crit", "Pa, –"],
         ["P_t, P_n", "Pas des tubes ; pas normal à l'écoulement", "m"],
         ["q″_w", "Densité de flux pariétal, rapportée à l'aire extérieure", "W/m²"],
         ["R_wall", "Résistance thermique nodale de paroi (cylindrique)", "K/W"],
         ["S_m, S_w, S_tb, S_sb, S_b", "Aires Bell-Delaware : transversale, fenêtre, tube-chicane, calandre-chicane, contournement", "m²"],
         ["s_w", "Épaisseur de paroi du tube", "m"],
         ["T_g, T_c", "Températures de mélange du gaz et du réfrigérant", "K"],
         ["T_wg, T_wc", "Températures de paroi côté gaz et côté réfrigérant", "K"],
         ["T_pc", "Température pseudo-critique à la pression considérée", "K"],
         ["UA", "Conductance globale nodale", "W/K"],
         ["U_g", "Vitesse du gaz", "m/s"],
         ["x", "Titre thermodynamique (fraction massique de vapeur)", "–"],
         ["X_tt", "Paramètre de Lockhart-Martinelli, turbulent-turbulent", "–"],
         ["α", "Taux de vide", "–"],
         ["Δx", "Longueur de maille axiale", "m"],
         ["μ, ρ, ν", "Viscosité, masse volumique, coefficient de Poisson", "Pa·s, kg/m³, –"],
         ["φ", "Indice de sévérité de corrugation, e²/(p D_i)", "–"],
         ["ω", "Facteur de sous-relaxation", "–"]],
        "Principaux symboles. Les indices g et c désignent respectivement le gaz chaud et le "
        "réfrigérant ; l et v le liquide et la vapeur saturés ; b et w les conditions de mélange "
        "et de paroi.",
        widths=[4.4, 9.4, 2.4], mono_cols=(0,))

    # ---------------- 2 ----------------
    D.h1("2.  Structure du problème")

    D.p("Les conditions aux limites déterminent à la fois la nature mathématique du problème et la "
        "stratégie de résolution. Les deux entrées de fluide sont imposées : le gaz entre à l'entrée "
        "du tube avec une température et une composition fixées par le calcul de combustion amont, et "
        "le réfrigérant entre à sa propre entrée avec une température et une pression imposées. Le "
        "point déterminant est la **position relative** de ces deux entrées, qui dépend de la "
        "configuration d'écoulement.")

    D.table(
        ["Configuration", "Entrée gaz", "Entrée réfrigérant", "Nature mathématique"],
        [["Co-courant", "nœud 0", "nœud 0",
          "**Problème de Cauchy.** Toutes les données sont disponibles à une même extrémité ; une "
          "unique marche avant simultanée suffirait en principe."],
         ["Contre-courant", "nœud 0", "nœud N−1",
          "**Problème aux limites à deux points.** Au nœud 0 la température du gaz est connue, mais "
          "pas celle du réfrigérant — c'est sa sortie, la seule grandeur non imposée."]],
        "Structure des conditions aux limites selon la configuration. C'est le contre-courant qui "
        "impose une résolution itérative ; la même voie est employée pour le co-courant plutôt que "
        "de maintenir un second chemin de code.",
        widths=[2.9, 2.0, 2.4, 8.9])

    D.p("Un problème aux limites à deux points ne peut être intégré par une simple marche avant, la "
        "moitié des données requises se trouvant à l'extrémité opposée du domaine. Le solveur "
        "décompose donc le problème en deux marches unidirectionnelles — une côté tubes, une côté "
        "calandre — et les itère l'une contre l'autre jusqu'à cohérence mutuelle. C'est le "
        "**balayage prédictif** de la section 8.")

    D.note("Pourquoi la décomposition est licite",
           "Chaque marche est, isolément, un problème de Cauchy bien posé. À champ de température "
           "de calandre figé, la marche côté tubes est une pure intégration avant ; à champ de "
           "puissance figé, la marche énergétique côté calandre l'est également. Seul leur couplage "
           "est implicite, et c'est précisément ce couplage que l'itération externe résout.")

    D.p("Une conséquence mérite d'être notée : les deux marches étant déjà découplées, le changement "
        "de sens d'écoulement est peu coûteux. Le contre-courant ne requiert que l'inversion de la "
        "boucle de marche côté calandre, ainsi qu'un suivi correct de la distance depuis l'entrée "
        "propre du réfrigérant pour les corrections de longueur d'établissement. Aucune méthode de "
        "tir ni recherche de racine n'est nécessaire. Cela diffère du solveur à serpentin hélicoïdal "
        "du même code, qui intègre paroi et réfrigérant en une seule passe couplée et doit donc "
        "résoudre le contre-courant par tir sur l'enthalpie de départ côté chaud, avec encadrement "
        "adaptatif et dichotomie.")

    # ---------------- 3 ----------------
    D.h1("3.  Discrétisation")

    D.p("Un maillage axial uniforme de N mailles couvre la longueur du tube, le nœud i étant centré "
        "en x = (i + ½)Δx :")
    D.eq(r"\Delta x = \frac{L_{tube}}{N} \qquad N = 200 \text{ (par défaut)}")

    D.p("L'indice 0 désigne toujours l'**extrémité d'entrée du gaz**, quel que soit le sens du "
        "réfrigérant. En co-courant le réfrigérant débute également à l'indice 0 ; en contre-courant "
        "il débute à l'indice N−1 et progresse en sens inverse. Conserver une convention d'indice "
        "unique pour les deux fluides permet à un seul tableau de puissance de servir les deux "
        "marches sans réindexation.")

    D.p("La géométrie côté tubes découle du diamètre extérieur et de l'épaisseur de paroi, et le "
        "débit de gaz est réparti sur le faisceau :")
    D.eqs(r"D_i = D_o - 2 s_w",
          r"A_i = \frac{\pi D_i^{2}}{4}",
          r"\dot{m}_{tube} = \frac{\dot{m}_g}{N_{tubes}}")

    D.note("Grandeurs par tube et grandeurs de faisceau",
           "Toute grandeur côté tubes de ce modèle — vitesse, nombre de Reynolds, bilan enthalpique "
           "et puissance nodale dQ_i — est définie **par tube représentatif**. La multiplication par "
           "N_tubes intervient exactement une fois, au passage de la puissance côté calandre "
           "(" + D.ref("ntubes") + "). Utiliser le débit de gaz total dans la vitesse côté tubes est "
           "une erreur fréquente et silencieuse.",
           colour=FLAG)

    D.note("S_sb corrigé le 2026-08-20",
           "L'aire de fuite calandre-chicane est le jeu diamétral sur l'arc où le bord de la "
           "chicane suit la calandre : S_sb = Lsb*(Ds/2)*(2*pi - theta_ds) = "
           "Ds*Lsb*(pi - theta_ds/2). L'implémentation portait auparavant un facteur 1/2 "
           "parasite. Confirmé par la dérivation élémentaire jeu x longueur d'arc et par "
           "Hellborg (2017) éq. 47, qui concordent. Effet sur cette géométrie : "
           "S_sb 1,218 -> 2,436 cm², r_lm 6,51 -> 6,93.",
           colour=FLAG)

    D.p("Côté calandre, la géométrie se réduit au jeu d'aires de Bell-Delaware. L'aire déterminante "
        "est l'aire de passage transversale minimale au plan médian de la calandre, qui fixe le flux "
        "massique de référence de toutes les corrélations côté calandre :")
    D.eq(r"S_m = B \left[ (D_s - D_{otl}) + \frac{D_{otl} - D_o}{P_n}(P_t - D_o) \right]")

    # ---------------- 4 ----------------
    D.h1("4.  Modèle côté tubes (gaz chaud)")

    D.p("Une marche avant du nœud 0 au nœud N−1, évaluée contre le champ de température de calandre "
        "courant.")

    D.h2("4.1  Paramétrage par enthalpie extraite")

    D.p("Plutôt que de transporter la température comme variable d'état, la marche transporte "
        "l'**enthalpie massique cumulée extraite** du gaz. Dans le mode cinétique à vitesse finie, "
        "qui est le mode par défaut, l'état thermochimique complet est lu dans une table "
        "flamelette/variable d'avancement (FPV) précalculée, indexée par deux scalaires : "
        "l'enthalpie extraite et une variable d'avancement de recombinaison non normalisée Y_c :")
    D.eq(r"(T_g, \rho_g, \mu_g, k_g, c_{p,g}, \omega_{Yc}) = M(h_{removed}, Y_c)")

    D.p("L'interpolation dans la table est bilinéaire. La variable d'avancement est transportée par "
        "son propre taux de production net, ce qui permet à la composition de retarder sur "
        "l'équilibre à mesure que le gaz se refroidit ; ce retard *est* la physique de la cinétique "
        "à vitesse finie :")
    D.eqs(r"\frac{dY_c}{dx} = \frac{\omega_{Yc}(h_{removed}, Y_c)}{U_g}",
          r"\frac{dh_{removed}}{dx} = \frac{1}{\dot{m}_{tube}} \frac{dQ}{dx}")

    D.p("Les deux sont avancées explicitement, nœud par nœud. Deux modes de comparaison remplacent "
        "cette lecture de table par un appel direct à un solveur d'équilibre chimique à chaque nœud, "
        "rééquilibrant le mélange à pression constante après extraction d'enthalpie (mode équilibre) "
        "ou figeant la composition (mode gelé). Le mode cinétique est le mode physiquement approprié "
        "au régime de forte extraction de chaleur considéré ; les deux autres ne servent qu'à la "
        "validation.")

    D.h2("4.2  Groupes adimensionnels et film côté gaz")

    D.eqs(r"U_g = \frac{\dot{m}_{tube}}{\rho_g A_i}",
          r"Re_g = \frac{\rho_g U_g D_i}{\mu_g}",
          r"Pr_g = \frac{c_{p,g}\, \mu_g}{k_g}",
          r"h_g = \frac{Nu_g\, k_g}{D_i}")

    D.p("Le nombre de Nusselt et le coefficient de frottement sont sélectionnés selon le traitement "
        "de surface interne. Pour les tubes corrugués, tous deux proviennent des corrélations de "
        "Vicente-Cruz, paramétrées par un indice de sévérité de corrugation construit sur la "
        "profondeur e et le pas p des cannelures :")
    D.eq(r"\varphi = \frac{e^{2}}{p\, D_i}")

    D.p("Les tubes lisses emploient la sélection standard pour tube droit avec frottement de "
        "Colebrook. Dans les deux cas, le résultat est mis à l'échelle par des multiplicateurs de "
        "calage contenus dans la structure `CorrelationCoefficients`, source unique des paramètres "
        "d'ajustement du modèle.")

    D.h2("4.3  Quantité de mouvement côté gaz")

    D.eq(r"\frac{dp_g}{dx} = - \frac{f_g \rho_g U_g^{2}}{2 D_i}")

    D.note("Convention sur le coefficient de frottement",
           "f_g est le coefficient de **Darcy** dans tous les chemins maintenus du solveur, en "
           "cohérence avec le facteur deux au dénominateur ci-dessus. Appliquer par-dessus une "
           "conversion Fanning-vers-Darcy quadruple la perte de charge prédite.",
           colour=FLAG)

    D.p("Le couplage est ici délibérément unidirectionnel : la pression du gaz est intégrée à titre "
        "de restitution, mais n'est pas réinjectée dans la masse volumique, la table FPV étant "
        "construite à une pression de référence fixe. Cela est admissible car le combusteur est "
        "dimensionné pour des nombres de Mach gaz inférieurs à environ 0,3, où l'erreur de masse "
        "volumique induite par la perte de charge reste faible. C'est également la raison pour "
        "laquelle aucune équation de quantité de mouvement quasi-1D n'est résolue pour le gaz chaud "
        "dans ce modèle.")

    # ---------------- 5 ----------------
    D.h1("5.  Modèle de conduction pariétale unidimensionnel")

    D.p("À chaque nœud, les deux films convectifs et la paroi du tube forment un réseau de trois "
        "résistances en série entre les températures de mélange du gaz et du réfrigérant. La "
        "particularité géométrique essentielle est que les deux films agissent sur des **aires "
        "différentes**, le tube étant cylindrique et le fluide chaud circulant à l'intérieur.")

    D.figure(os.path.join(fig_dir, "fig_wall_network.png"), 16.2,
             "Disposition radiale et réseau de résistances nodal équivalent. Le fluide chaud étant "
             "intérieur, le film côté gaz agit sur le périmètre intérieur, plus petit, et le film côté "
             "réfrigérant sur le périmètre extérieur, plus grand.")

    D.h2("5.1  Géométrie et affectation du côté chaud")

    D.eqs(r"P_{hot} = \pi D_i \qquad P_{cold} = \pi D_o",
          r"A_{hot} = P_{hot}\, \Delta x \qquad A_{cold} = P_{cold}\, \Delta x")

    D.note("Orientation du côté chaud",
           "Le modèle de paroi est partagé avec la configuration à serpentin hélicoïdal, dans "
           "laquelle le fluide chaud est **extérieur** au tube et l'affectation des périmètres est "
           "inversée. L'orientation est choisie par un indicateur `hot_side`, qui doit valoir "
           "\"inner\" pour la configuration calandre-et-tubes. Un réglage erroné laisse le réseau "
           "structurellement intact mais met à l'échelle les deux résistances de film par D_o/D_i "
           "dans le mauvais sens, produisant une puissance plausible mais fausse.",
           colour=FLAG)

    D.h2("5.2  Résistance de paroi et conductance globale")

    D.p("La résistance de paroi est **cylindrique**, et non plane — distinction importante aux "
        "rapports épaisseur/diamètre employés ici :")
    D.eq(r"R_{wall} = \frac{\ln\left[(D_i/2 + s_w)/(D_i/2)\right]}{2 \pi \Delta x\, k_w}")

    D.p("Les trois résistances se composent en série pour donner la conductance nodale et la "
        "puissance échangée au nœud :")
    D.eqs(r"\frac{1}{UA} = \frac{1}{h_g A_{hot}} + R_{wall} + \frac{1}{h_c A_{cold}}",
          r"dQ_i = UA \left( T_{g,i} - T_{c,i} \right)")

    D.h2("5.3  Températures de face")

    D.p("Les températures de face de paroi s'obtiennent en parcourant la chaîne de résistances, la "
        "troisième relation servant de vérification de cohérence : la chaîne doit ramener à la "
        "température de réfrigérant imposée.")
    D.eqs(r"T_{wg} = T_g - \frac{dQ}{h_g A_{hot}}",
          r"T_{wc} = T_{wg} - dQ\, R_{wall}",
          r"T_c^{check} = T_{wc} - \frac{dQ}{h_c A_{cold}} \;\;\equiv\;\; T_c")

    D.h2("5.4  Conductivité dépendant de la température et résolution nodale")

    D.p("La conductivité de paroi est évaluée à la température moyenne dans l'épaisseur :")
    D.eq(r"k_w = k_w(\bar{T}_w), \qquad \bar{T}_w = \frac{T_{wg} + T_{wc}}{2}")

    D.p("La résistance de paroi dépend donc des températures de face que l'on cherche précisément à "
        "déterminer. Chaque nœud referme en conséquence un petit système non linéaire, résolu à une "
        "tolérance de 10⁻⁸ sur le vecteur résidu :")
    D.eq(r"F = \left[ T_{wg} - T_{wg}^{new},\; T_{wc} - T_{wc}^{new},\; T_c^{check} - T_c^{check,new} \right] = 0")

    D.p("Cette résolution non linéaire nodale est imbriquée dans la boucle sur les nœuds axiaux, "
        "elle-même imbriquée dans le balayage externe de la section 8 — soit trois niveaux "
        "d'itération au total.")

    D.h2("5.5  Flux pariétal et rayonnement")

    D.p("La densité de flux pariétal transmise aux fermetures côté calandre est rapportée à l'aire "
        "**extérieure**, puisque c'est la surface que voit le réfrigérant :")
    D.eq(r"q''_{w,i} = \frac{dQ_i}{\pi D_o \Delta x}")

    D.p("Un modèle de rayonnement du gaz (émissivité par longueur de faisceau moyenne, intervenant "
        "comme coefficient radiatif en parallèle du film convectif côté gaz) est implémenté dans le "
        "module de paroi mais **désactivé** dans le chemin stationnaire calandre-et-tubes.")

    # ---------------- 6 ----------------
    D.h1("6.  Échange thermique et perte de charge côté calandre")

    D.h2("6.1  Niveau de description")

    D.p("L'écoulement côté calandre traverse le faisceau transversalement, contourne une chicane, "
        "puis traverse de nouveau, en alternance sur la longueur de l'échangeur. Sur l'essentiel de "
        "son trajet, le vecteur vitesse est perpendiculaire à la coordonnée axiale, et il n'existe "
        "aucune coordonnée curviligne unique le long de laquelle poser une équation de quantité de "
        "mouvement unidimensionnelle. Le côté calandre est donc modélisé comme **zéro-dimensionnel "
        "par nœud axial** : une corrélation globale de faisceau évaluée avec les propriétés locales. "
        "Il ne s'agit pas d'une simplification d'une équation différentielle de quantité de "
        "mouvement, mais d'une forme de modèle différente et plus appropriée, calée empiriquement "
        "sur cette géométrie précise.")

    D.h2("6.2  Faisceau idéal de Bell-Delaware")

    D.eqs(r"G_s = \frac{\dot{m}_c}{S_m}",
          r"Re_s = \frac{D_o G_s}{\mu_s}",
          r"Pr_s = \frac{c_p \mu}{k}")

    D.p("Le facteur j de Colburn du faisceau idéal emploie des coefficients a₁–a₄ tabulés par "
        "arrangement et par plage de Reynolds, pour les dispositions triangulaire, carrée et carrée "
        "tournée :")
    D.eqs(r"a = \frac{a_3}{1 + 0.14\, Re_s^{a_4}}",
          r"j = a_1 \left( \frac{1.33}{P_t/D_o} \right)^{a} Re_s^{a_2}",
          r"h_{ideal} = j\, c_p\, G_s\, Pr_s^{-2/3} \left( \frac{\mu_b}{\mu_w} \right)^{0.14}")

    D.note("La correction de propriétés est évaluée, non laissée à sa valeur neutre",
           "Le terme de Sieder-Tate (mu_b/mu_w)^0.14 ci-dessus est évalué à partir de la température "
           "de paroi côté réfrigérant retardée (section 8.1). Il était auparavant laissé à sa valeur "
           "neutre de 1,0, c'est-à-dire que la correction de variation de propriétés propre à la "
           "corrélation était désactivée. Son activation a abaissé de 67 K la température de paroi "
           "maximale du point de fonctionnement Eau et résolu un blocage de convergence sur ce point ; "
           "elle déplace la référence Hélium calandre-et-tubes d'environ 0,1 %. Le rapport est borné "
           "à [0,25 ; 4] afin qu'un état de paroi non convergé en début de calcul ne puisse fausser "
           "le coefficient.")

    D.p("Cinq facteurs correctifs multiplicatifs rendent ensuite compte des écarts du faisceau réel "
        "au faisceau idéal :")
    D.eq(r"h_{shell} = h_{ideal}\, J_c\, J_l\, J_b\, J_s\, J_r", label="hshell")

    D.table(
        ["Facteur", "Effet physique", "Forme"],
        [["J_c", "Configuration de fenêtre de chicane : les tubes situés dans la fenêtre voient un "
                 "écoulement axial et non transversal. F_c est la fraction de tubes en écoulement "
                 "purement transversal.",
          "0,55 + 0,72 F_c"],
         ["J_l", "Fuites aux chicanes par les jeux tube-chicane et calandre-chicane, qui "
                 "court-circuitent le faisceau. r_lm = (S_sb+S_tb)/S_m, r_s = S_sb/(S_sb+S_tb).",
          "p + (1−p)exp(−2,2 r_lm),  p = 0,44(1−r_s)"],
         ["J_b", "Contournement du faisceau par sa périphérie. C = 1,35 en laminaire, 1,25 en "
                 "turbulent ; vaut 1 lorsque le rapport de bandes d'étanchéité r_ss ≥ 0,5.",
          "exp(−C F_sbp [1−(2 r_ss)^(1/3)])"],
         ["J_s", "Espacements de chicane d'entrée et de sortie différents de l'espacement central.",
          "rapport d'espacements, exposant n₁ = 0,6"],
         ["J_r", "Gradient de température défavorable en laminaire. Actif seulement en deçà de "
                 "Re ≈ 100, vaut 1 au-delà.",
          "(10/N_c)^0,18, raccordé linéairement à 1 en Re = 100"]],
        "Facteurs correctifs de Bell-Delaware. Ordres de grandeur typiques pour une coupe de chicane "
        "de 25 % : J_c ≈ 1,0, J_l ≈ 0,7–0,8, J_b ≈ 0,9, J_s ≈ 1 et J_r = 1 en turbulent.",
        widths=[1.6, 9.0, 5.8], mono_cols=(0,))

    D.h2("6.3  Perte de charge côté calandre")

    D.p("La perte de charge est assemblée à partir de trois zones — les passages transversaux "
        "centraux, les fenêtres de chicane et les deux zones d'extrémité — chacune portant ses "
        "propres corrections de fuite et de contournement (R_l, R_b, R_s) :")
    D.eqs(r"\Delta p_{ideal} = 2 f \frac{G_s^{2}}{\rho} N_{tcc} \left( \frac{\mu_b}{\mu_w} \right)^{-0.14}",
          r"\Delta p_{cross} = \Delta p_{ideal} (N_b - 1) R_b R_l",
          r"\Delta p_{window} = (2 + 0.6 N_{tcw}) \frac{G_w^{2}}{2\rho} N_b R_l, \qquad G_w = \frac{\dot{m}_c}{\sqrt{S_m S_w}}",
          r"\Delta p_{ends} = \Delta p_{ideal} \left( 1 + \frac{N_{tcw}}{N_{tcc}} \right) R_b R_s",
          r"\Delta p_{shell} = \Delta p_{cross} + \Delta p_{window} + \Delta p_{ends}",
          label="dpshell")

    D.h3("6.3.1  Perte de charge diphasique")

    D.p("À l'intérieur du dôme, la perte de charge est rapportée à la valeur Bell-Delaware "
        "TOUT LIQUIDE puis mise à l'échelle par le multiplicateur diphasique de Chisholm, "
        "plutôt que d'évaluer Bell-Delaware à la masse volumique homogène du mélange. Le "
        "multiplicateur restitue exactement les deux limites — l'unité en x=0 et Γ² en x=1 :")
    D.eqs(r"\Gamma = \sqrt{\frac{\rho_l}{\rho_v}} \left( \frac{\mu_v}{\mu_l} \right)^{0.1}",
          r"\phi^{2} = 1 + (\Gamma^{2} - 1) \left[ B\, x^{(2-n)/2} (1-x)^{(2-n)/2} + x^{2-n} \right]",
          r"\Delta p_{TP} = \phi^{2}\, \Delta p_{liquide}",
          label="chisholm")
    D.p("avec n = 0,2 en régime turbulent. B est sélectionné selon (Γ, flux massique) ; au "
        "point de fonctionnement Eau de cet échangeur, Γ = 3,60 et G ≈ 3000 kg/m²·s, "
        "d'où B ≈ 1,01.")

    D.note("B est reconstitué, non transcrit",
           "Chisholm (1973) et Grant & Chisholm (1979) n'étaient pas disponibles lors de "
           "l'implémentation. La table de B n'est corroborée qu'en un seul point : Hellborg "
           "(2017) éq. 137 code en dur B = 21/Γ, valeur qui tombe précisément dans la "
           "branche 9,5 < Γ < 28. C'est encourageant, ce n'est pas une vérification. Noter "
           "en outre que la branche codée en dur par Hellborg est FAUSSE pour cet "
           "échangeur : à Γ = 3,60 la branche correcte donne B ≈ 1,01 et non 5,83, de sorte "
           "que recopier cette équation telle quelle gonflerait le multiplicateur d'un "
           "facteur six environ.",
           colour=FLAG)

    D.h3("6.3.2  Perte de charge par accélération")

    D.p("À mesure que le réfrigérant se vaporise, sa masse volumique s'effondre et "
        "l'écoulement doit accélérer, ce qui coûte une pression supplémentaire au-delà du "
        "frottement pariétal. Sur une maille :")
    D.eq(r"\Delta p_{acc} = G^{2} \left[ \left( \frac{1}{\rho} \right)_{i+1} - \left( \frac{1}{\rho} \right)_{i} \right]",
         label="dpacc")
    D.p("Il s'agit de la forme homogène : le terme de quantité de mouvement en écoulement "
        "séparé x²/(α·ρ_v) + (1−x)²/((1−α)·ρ_l) se réduit à 1/ρ pour le taux de vide "
        "homogène que restitue la fermeture d'état ; aller au-delà exigerait un α issu d'un "
        "modèle à dérive. Ce terme était purement et simplement omis ; il représente "
        "26 % de la perte de charge totale côté calandre au point Eau (2,99 bar sur 11,05, "
        "la masse volumique passant de 1000 à 29 kg/m³) et 23 % au point N₂.")

    D.note("Un jeu de corrélations solidaire",
           "Les tables de coefficients j et f du faisceau idéal et les facteurs correctifs J/R ont "
           "été calés les uns contre les autres. Substituer une autre corrélation de faisceau sans "
           "redériver les corrections invalide le calage, même si chaque composant est "
           "individuellement défendable.")

    D.h2("6.4  Sélection du régime")

    D.p("Bell-Delaware ne contient aucune physique de l'ébullition : à l'intérieur du dôme "
        "diphasique, elle doit donc être abandonnée purement et simplement. À pression "
        "supercritique, la situation est plus subtile : ce que corrigent les fermetures à rapport de "
        "propriétés, c'est la forte variation des propriétés au voisinage de la température "
        "pseudo-critique T_pc(p) — vestige étalé du pic de chaleur latente. Loin de cette région, un "
        "fluide supercritique est un fluide monophasique ordinaire, et la corrélation de "
        "Bell-Delaware, calée sur l'écoulement transversal, est le modèle le mieux adapté à un "
        "faisceau chicané. La sélection s'opère donc sur l'état local :")

    D.table(
        ["Régime local du réfrigérant", "Fermeture d'échange thermique", "Fermeture de frottement"],
        [["Liquide sous-refroidi ou vapeur surchauffée (monophasique, sous-critique)",
          "Bell-Delaware, " + D.ref("hshell"), "Bell-Delaware, " + D.ref("dpshell")],
         ["Diphasique saturé, 0 ≤ x ≤ 1",
          "Gungor-Winterton, " + D.ref("gw"), "Müller-Steinhagen & Heck, " + D.ref("msh")],
         ["Supercritique, intervalle mélange-paroi atteignant la bande pseudo-critique autour de T_pc(p)",
          "Fermeture à rapport de propriétés (McCarthy-Wolf / Taylor)",
          "Gradient de frottement de la fermeture"],
         ["Supercritique, mais loin de T_pc — aucune correction de propriétés justifiée",
          "Bell-Delaware, " + D.ref("hshell"),
          "Gradient de frottement de la fermeture"]],
        "Sélection de la fermeture côté calandre. Le régime est évalué par nœud et par balayage, à "
        "partir de l'état (p, h) courant du réfrigérant et de la température de paroi retardée. "
        "Noter la dernière ligne : la fermeture d'échange thermique bascule vers Bell-Delaware "
        "tandis que le modèle de frottement, lui, ne bascule pas — voir ci-dessous.",
        widths=[6.4, 6.2, 3.8])

    D.note("La pression supercritique n'est pas à elle seule le critère",
           "La sélection était auparavant conditionnée uniquement par la chaîne de caractères "
           "coolantProp.coolant_model : la fermeture attribuée à un fluide relevait donc d'un "
           "artefact de configuration plutôt que d'un constat sur son état. Hélium et N₂ sont tous "
           "deux supercritiques mais se situent en des points entièrement différents : l'Hélium à "
           "80 bar a T_pc = 11,4 K face à une marche de 300 à 1400 K — soit 26 à 120 fois au-dessus "
           "de toute anomalie critique, avec un cp plat à 0,1 % près — tandis que le N₂ à 88 bar a "
           "T_pc = 147,8 K, un mélange à 100-124 K et une paroi atteignant 164 K, ce qui place la "
           "transition pseudo-critique **à l'intérieur de la couche limite thermique**. Au point de "
           "fonctionnement N₂, 97 nœuds sur 200 satisfont le critère. La sélection est verrouillée "
           "par nœud, de façon irréversible, afin qu'elle ne puisse osciller pendant que la "
           "température de paroi retardée s'élève encore depuis son amorce froide.",
           colour=FLAG)

    D.note("La marche en pression reste sur un seul modèle de frottement",
           "Bell-Delaware restitue une perte de charge **globale** de faisceau, répartie en "
           "dp_shell/N, alors que la fermeture restitue un **gradient local**. Mélanger les deux "
           "d'un nœud à l'autre revient à compter deux fois, chaque nœud Bell-Delaware apportant une "
           "quote-part d'une perte de charge globale que les autres nœuds ne prennent pas en charge. "
           "Un nœud supercritique qui bascule vers Bell-Delaware pour son coefficient d'échange "
           "conserve donc un gradient de pression issu du modèle local. Sur cette géométrie, la "
           "perte de charge de Bell-Delaware est de toute façon inutilisable — voir section 11.")

    D.h2("6.5  Ébullition convective saturée")

    D.p("À l'intérieur du dôme diphasique, la corrélation de Gungor-Winterton (1986) est employée. "
        "Sa structure à deux termes porte la physique : la convection forcée est amplifiée par "
        "l'accélération de la vapeur, tandis que l'ébullition nucléée est supprimée à mesure que "
        "cette convection amincit la couche surchauffée pariétale.")
    D.eqs(r"Re_l = \frac{G(1-x)D}{\mu_l}, \qquad h_l = 0.023\, Re_l^{0.8} Pr_l^{0.4} \frac{k_l}{D}",
          r"X_{tt} = \left( \frac{1-x}{x} \right)^{0.9} \left( \frac{\rho_v}{\rho_l} \right)^{0.5} \left( \frac{\mu_l}{\mu_v} \right)^{0.1}",
          r"Bo = \frac{q''}{G\, h_{fg}}",
          r"E = 1 + 24000\, Bo^{1.16} + 1.37 \left( \frac{1}{X_{tt}} \right)^{0.86}",
          r"S = \frac{1}{1 + 1.15 \cdot 10^{-6} E^{2} Re_l^{1.17}}",
          r"h_{pool} = 55\, p_r^{0.12} (-\log_{10} p_r)^{-0.55} M^{-0.5} q''^{0.67}",
          r"h = E\, h_l + S\, h_{pool}",
          label="gw")

    D.p("Le terme d'ébullition en vase est la corrélation de Cooper, M désignant la masse molaire en "
        "kg/kmol. La correction pour faible nombre de Froude en tube horizontal, présente dans la "
        "référence originale, n'est pas appliquée : la corrélation est employée ici sous sa forme "
        "verticale / à Froude élevé.")

    D.h2("6.6  Frottement diphasique")

    D.p("La corrélation de Müller-Steinhagen et Heck interpole empiriquement entre les limites tout "
        "liquide et tout vapeur, en restituant correctement chacune en x = 0 et x = 1 :")
    D.eqs(r"A = \frac{f_{l0} G^{2}}{2 D \rho_l}, \qquad B = \frac{f_{v0} G^{2}}{2 D \rho_v}",
          r"-\frac{dp}{dz} = \left[ A + 2(B-A)x \right] (1-x)^{1/3} + B x^{3}",
          label="msh")

    D.h2("6.7  Flux critique et déclenchement de l'ébullition nucléée")

    D.p("L'écart à l'ébullition nucléée est surveillé au moyen d'un flux thermique critique tabulé "
        "(Groeneveld 2006), restitué sous forme de marge :")
    D.eq(r"\text{marge CHF} = \frac{q''_{CHF}}{q''_w}")

    D.note("Non définie hors du dôme",
           "La marge de flux critique est restituée comme NaN pour les états liquide sous-refroidi, "
           "vapeur surchauffée et supercritique. Il s'agit d'une véritable frontière de régime — la "
           "notion de flux thermique critique n'y existe pas — et non d'une donnée manquante. Une "
           "marge qui s'annule exactement en x → 1 traduit la transition d'assèchement vers la "
           "vaporisation complète, comportement recherché pour un dimensionnement qui surchauffe "
           "délibérément son réfrigérant.")

    D.p("Deux mécanismes distincts encadrent la transition sous-refroidi/ébullition et ne doivent "
        "pas être confondus. Le premier est une **fenêtre de raccordement numérique** de demi-largeur "
        "0,02 en titre, centrée sur x = 0, qui lisse une variation autrement abrupte du coefficient "
        "d'échange entre les branches monophasique et diphasique. C'est un dispositif de lissage sans "
        "contenu physique. Le second est le critère **physique** de Bergles-Rohsenow (1964), "
        "surchauffe pariétale de déclenchement de l'ébullition nucléée, évaluée comme diagnostic et "
        "signalée lorsqu'elle est dépassée :")
    D.eq(r"q''_{ONB} = 1082\, p^{1.156} \left( 1.8\, \Delta T_{ONB} \right)^{2.16/p^{0.0234}}")

    D.p("avec p en bar, inversée pour ΔT_ONB et comparée à la surchauffe pariétale estimée. "
        "L'ébullition nucléée sous-refroidie peut débuter bien avant que le fluide de mélange "
        "n'atteigne la saturation si la paroi est suffisamment chaude ; lorsque le solveur signale "
        "un dépassement ONB, il indique que le raccordement numérique lisse tient lieu d'une "
        "ébullition que la physique dit déjà amorcée.")

    # ---------------- 7 ----------------
    D.h1("7.  Marche énergétique côté calandre")

    D.p("La puissance nodale calculée par tube représentatif est portée au faisceau complet une "
        "seule fois, ici :")
    D.eq(r"dQ_{total,i} = dQ_i\, N_{tubes}", label="ntubes")

    D.h2("7.1  Marche monophasique (en température)")
    D.eq(r"T_{i+1} = T_i + \frac{dQ_{total,i}}{\dot{m}_c\, c_{p,c}(T_i)}", label="Tmarch")

    D.h2("7.2  Marche fluide réel (en enthalpie)")
    D.eqs(r"h_{i+1} = h_i + \frac{dQ_{total,i}}{\dot{m}_c}",
          r"p_{i+1} = \max \left( p_i - \Delta p_i^{lagged},\; 1 \right)",
          r"(T, x, \alpha)_i = \text{flash}_{ph}(p_i, h_i)",
          label="hmarch")

    D.note("Pourquoi l'enthalpie plutôt que la température",
           "Une mise à jour en cp·ΔT n'a nulle part où placer la chaleur latente. Intégré en "
           "température, un réfrigérant en ébullition continuerait de monter au-delà de 1000 K au "
           "lieu de se bloquer près de la saturation le temps de se vaporiser. La distinction de "
           "fond porte sur la variable qui détermine l'état thermodynamique : en écoulement "
           "monophasique à pression fixée, h = h(T) est monotone et inversible, de sorte que les "
           "deux marches portent la même information. À l'intérieur du dôme, T = T_sat(p) est "
           "**dégénérée** — identique en x = 0,1 et en x = 0,9 — si bien que la température ne peut "
           "plus représenter l'état, alors que le couple (p, h) le peut encore.",
           colour=FLAG)

    D.p("La marche monophasique est une approximation d'Euler explicite de la marche en enthalpie, "
        "cp étant évalué à la température de début de maille. Elle est exacte à cp constant et "
        "d'ordre un sinon, l'erreur décroissant en 1/N. Pour l'Hélium, monoatomique et dont le cp est "
        "plat à environ 0,1 % près entre 300 K et 1400 K sous 80 bar, l'erreur sur la température de "
        "sortie qui en résulte est indécelable. La section 10 quantifie ce point par fluide.")

    D.h2("7.3  Marche en pression et retard d'une itération")

    D.p("La perte de charge nodale employée en " + D.ref("hmarch") + " provient du balayage "
        "**précédent**. Ce retard est inévitable : le profil de pression dépend du profil "
        "d'enthalpie par la masse volumique et le titre locaux, tandis que le profil d'enthalpie "
        "dépend du profil de pression par la température de saturation locale. Retarder l'un des "
        "deux brise la dépendance circulaire à l'ordre un dans l'itération de balayage, et le retard "
        "s'annule à convergence.")

    D.h2("7.4  Traitement de la limite de validité de l'équation d'état")

    D.p("Un débit de réfrigérant très faible face à une forte puissance peut surchauffer le fluide "
        "au-delà du domaine de validité de la bibliothèque de propriétés. Il s'agit d'un cas "
        "physique limite réel et non d'une défaillance numérique. En cas d'échec du calcul d'état, "
        "le dernier triplet (état, enthalpie, pression) valide est figé pour le reste de la marche. "
        "Les trois doivent être figés ensemble : ne figer que l'état en laissant dériver enthalpie "
        "et pression provoquerait le même échec au balayage suivant, lorsque ces valeurs dérivées "
        "sont relues.")

    # ---------------- 8 ----------------
    D.h1("8.  Itération externe (balayage prédictif)")

    D.figure(os.path.join(fig_dir, "fig_sweep.png"), 16.2,
             "Le balayage prédictif. Les deux marches échangent un champ de puissance vers l'aval et "
             "un champ de température vers l'amont. Trois entrées de fermeture ne peuvent être "
             "évaluées avant la marche qui les produit et sont donc consommées avec un balayage de "
             "retard.")

    D.p("Les deux marches sont refermées par une itération de point fixe (Picard) sous-relaxée :")
    D.eqs(r"tube \leftarrow \text{march}(T_{shell}^{(k)})",
          r"T_{new} \leftarrow \text{shell march}(dQ)",
          r"\delta = \max \left| T_{new} - T_{shell}^{(k)} \right|",
          r"T_{shell}^{(k+1)} = (1-\omega) T_{shell}^{(k)} + \omega\, T_{new}")

    D.p("L'itération se poursuit jusqu'à δ < 0,05 K, dans la limite de 25 balayages. La "
        "sous-relaxation à ω = 0,5 est fonctionnellement nécessaire et non cosmétique : le couplage "
        "est raide, le coefficient d'échange côté réfrigérant dépendant de la température de "
        "calandre, qui dépend de la puissance, qui dépend à son tour du coefficient d'échange. Une "
        "mise à jour non relaxée oscille ou diverge, en particulier au passage de la transition "
        "d'ébullition où la sensibilité du coefficient d'échange à la température de calandre est "
        "maximale.")

    D.h2("8.1  Grandeurs de fermeture retardées")

    D.table(
        ["Grandeur", "Consommée par", "Amorcée à", "Raison du retard"],
        [["q″_w", "Terme en nombre d'ébullition de Gungor-Winterton", "0",
          "Le flux pariétal est une sortie de la marche côté tubes dont la fermeture côté calandre a besoin en entrée"],
         ["Δp par nœud", "Marche en pression côté calandre", "0",
          "Les profils de pression et d'enthalpie sont mutuellement dépendants"],
         ["T_wc", "Rapport de propriétés des fermetures supercritiques ; terme de Sieder-Tate de Bell-Delaware",
          "T d'entrée réfrigérant",
          "La température de paroi côté réfrigérant est une sortie de la résolution de paroi"]],
        "Grandeurs retardées d'un balayage. Toutes sont amorcées de façon neutre et convergent vers "
        "la cohérence lorsque δ → 0.",
        widths=[2.6, 5.4, 2.6, 5.8], mono_cols=(0,))

    D.h2("8.2  Initialisation")

    D.note("L'amorçage diffère selon le modèle de réfrigérant",
           "Le mode monophasique peut amorcer le champ de calandre par un profil linéaire entre les "
           "deux entrées imposées. Le mode fluide réel ne le peut pas : une température de sortie "
           "réfrigérant imposée n'a pas de sens lorsque le fluide peut être diphasique, car à "
           "l'intérieur du dôme T et p ne sont pas indépendantes et un unique couple (T, p) ne peut "
           "représenter l'état. Le mode fluide réel amorce donc uniformément à la condition d'entrée "
           "sous-refroidie physique.")

    D.h2("8.3  Comportement en convergence")

    D.p("La convergence est monotone pour les cas monophasiques bien conditionnés. Au voisinage "
        "d'une transition d'ébullition, l'itération peut au contraire s'installer dans un cycle "
        "limite d'amplitude de l'ordre de 0,5 à 1 K sans amener δ sous la tolérance. La valeur de "
        "sortie restituée est stable à l'intérieur de cette bande, mais il s'agit d'une véritable "
        "non-convergence, signalée comme telle plutôt qu'acceptée silencieusement. Lorsqu'un point "
        "de fonctionnement présente ce comportement, il convient de le consigner avec le résultat.")

    D.h2("8.4  Bilan énergétique global")

    D.eq(r"Q_{tot} = \sum_{i}{ dQ_i\, N_{tubes} }")

    D.p("Le côté gaz est toujours intégré en enthalpie : la puissance retranchée à son bilan est "
        "donc exactement celle qui traverse la paroi. Que le côté réfrigérant se referme exactement "
        "dépend de la marche de la section 7 employée :")
    D.bullet("**Mode fluide réel (enthalpie) — exact par construction.** L'enthalpie étant la "
             "variable intégrée, le réfrigérant absorbe précisément la puissance sommée, par "
             "définition. Aucune erreur de fermeture n'existe à surveiller.")
    D.bullet("**Mode monophasique (température) — d'ordre un.** La température est accumulée avec "
             "cp évalué en début de maille : l'élévation d'enthalpie vraie et la puissance sommée "
             "diffèrent donc en général. L'écart est nul à cp constant et décroît en 1/N.")
    D.p("Dans aucun des deux cas il n'y a de dérive d'intégration temporelle à accumuler, ce qui "
        "explique que le solveur stationnaire ne porte pas de diagnostic de résidu énergétique, "
        "contrairement aux solveurs transitoires où un tel diagnostic est indispensable.")

    # ---------------- 9 ----------------
    D.h1("9.  Vérification mécanique")

    D.p("Le chargement des tubes dans cette configuration est l'inverse de celui d'un canal "
        "régénératif : le réfrigérant haute pression entoure les tubes tandis que le gaz, à pression "
        "comparativement basse, circule à l'intérieur. Les tubes sont donc sous pression "
        "**externe** et le mode de ruine considéré est le flambage élastique et non l'éclatement. "
        "Cette vérification est appelée séparément du calcul thermique.")
    D.eqs(r"\Delta P_{ext} = p_{c,in} - p_0",
          r"\sigma_{ext} = \sigma_{ext}(\Delta P_{ext}, s_w, D_i)",
          r"P_{cr} = P_{cr}(E(\bar{T}), s_w, D_i, \nu)",
          r"\text{marge de flambage} = \frac{\left| \Delta P_{ext} \right|}{P_{cr}} \; < 1 \;\text{ pour la sécurité}")

    D.p("Le module d'Young est évalué à la moyenne de la température de paroi côté gaz maximale et "
        "de la température de paroi côté réfrigérant minimale, converties en degrés Celsius, les "
        "tables de propriétés matériaux étant en °C alors que le solveur est par ailleurs "
        "entièrement en unités SI.")

    D.note("Extrapolation au-delà du domaine caractérisé",
           "Les données de l'Inconel 718 couvrent −240 °C à 760 °C (1033,15 K). Au-delà, la courbe "
           "de limite d'élasticité est écrêtée à plat plutôt qu'extrapolée par une loi matériau "
           "physique : un rapport de contrainte calculé à 1100 K, par exemple, est arithmétiquement "
           "bien formé mais physiquement non étayé. La température de paroi maximale doit toujours "
           "être restituée à côté de la marge de flambage afin que cette condition reste visible.",
           colour=FLAG)

    # ---------------- 10 ----------------
    D.h1("10.  Application à l'Hélium, au LN₂/N₂ et à l'Eau")

    D.p("Les trois réfrigérants considérés sollicitent des branches différentes du modèle. La "
        "structure côté gaz, la paroi et l'itération sont identiques dans les trois cas ; seules la "
        "fermeture côté calandre et la variable d'intégration énergétique diffèrent. Cette section "
        "fait le lien entre le cadre exposé ci-dessus et les résultats restitués.")

    D.table(
        ["Réfrigérant", "Régime typique au point de fonctionnement", "Fermeture côté calandre",
         "Marche énergétique", "Statut"],
        [["Hélium",
          "Monophasique sur toute la longueur, très au-dessus de la pression critique mais au "
          "comportement de type gaz (T_pc = 11,4 K face à une marche de 300 à 1400 K)",
          "Bell-Delaware",
          D.ref("Tmarch"),
          "**Référence.** Le cp est plat à environ 0,1 % près sur la plage de fonctionnement : la "
          "marche en température est numériquement indiscernable d'une marche en enthalpie."],
         ["LN₂ / N₂ supercritique",
          "Supercritique sur toute la longueur, p_crit ≈ 34 bar. Titre non défini ; fluide dense de "
          "type liquide. T_pc = 147,8 K, mélange à 100-124 K, paroi à 164 K.",
          "Fermeture à rapport de propriétés (McCarthy-Wolf / Taylor) là où l'intervalle "
          "mélange-paroi atteint la bande pseudo-critique ; Bell-Delaware ailleurs",
          D.ref("hmarch"),
          "**Sensibilité bornée.** Les corrélations supercritiques restent des calages en tube droit "
          "appliqués à un écoulement transversal, mais substituer Bell-Delaware ne déplace Q_tot que "
          "de 1 % (le film réfrigérant ne représente que 0,58 % de la résistance totale) ; la "
          "température de paroi porte environ 54 K d'incertitude."],
         ["Eau",
          "Traverse le liquide sous-refroidi, le dôme diphasique et la vapeur surchauffée au sein "
          "d'un même échangeur lorsque le dimensionnement vise une sortie vapeur.",
          "Bell-Delaware → Gungor-Winterton → Bell-Delaware, avec raccordement à l'entrée du dôme",
          D.ref("hmarch"),
          "**Partiellement validé.** Les corrélations d'ébullition sont validées dans la littérature "
          "pour l'écoulement en tube ; l'intensification de l'ébullition par l'écoulement transversal "
          "est réelle et non modélisée."]],
        "Chemins de modèle par réfrigérant. La réserve commune aux deux cas fluide réel est que "
        "Gungor-Winterton, Müller-Steinhagen-Heck et les fermetures supercritiques à rapport de "
        "propriétés sont toutes des corrélations d'écoulement en tube appliquées à un écoulement "
        "transversal côté calandre.",
        widths=[2.2, 4.4, 3.4, 2.2, 4.2])

    D.note("Intensification de l'ébullition en écoulement transversal",
           "L'écoulement transversal côté calandre autour d'un faisceau intensifie l'échange en "
           "ébullition par rapport à un écoulement interne en tube, par des mécanismes qu'aucune des "
           "corrélations employées ici ne représente. Le modèle est donc attendu **conservatif** sur "
           "le coefficient d'échange côté réfrigérant dans les branches diphasique et supercritique. "
           "Il s'agit d'un point ouvert connu, et non d'un oubli de modélisation ; il est signalé "
           "par le solveur à l'exécution plutôt que laissé silencieux.",
           colour=FLAG)

    # ---------------- 11 ----------------
    D.h1("11.  Hypothèses et domaine de validité")

    D.table(
        ["Hypothèse", "Justification", "Où elle devient significative"],
        [["Le rapport de fuite aux chicanes r_lm de Bell-Delaware est très au-delà du domaine de calage",
          "r_lm = (S_sb+S_tb)/S_m = 6,5 sur cette géométrie, contre un domaine de calage de l'ordre "
          "de r_lm ≤ 1 : l'aire de jeu tube-chicane (17,7 cm²) vaut 6,5 fois l'aire de passage "
          "transversale (2,9 cm²)",
          "**Les deux corrections de fuite sont extrapolées** : J_l = 0,41 sur l'échange thermique et "
          "R_l = 0,009 sur la perte de charge. R_l écrase les termes transversal et de fenêtre, "
          "laissant le terme d'extrémité non corrigé dominer une perte de charge globale d'environ "
          "288 bar. La perte de charge côté calandre issue de Bell-Delaware est **inutilisable** sur "
          "cette géométrie ; les pressions de sortie restituées proviennent du modèle de gradient "
          "local. Concerne tout résultat Bell-Delaware, y compris la référence Hélium. Signalé à "
          "l'exécution."],
         ["Les fermetures diphasique et supercritique sont des corrélations en tube appliquées à un "
          "écoulement transversal",
          "Aucune fermeture d'ébullition ou supercritique validée en écoulement transversal côté "
          "calandre n'est disponible",
          "Tout cas diphasique côté calandre. Pour le point supercritique N₂, la sensibilité a été "
          "mesurée : substituer Bell-Delaware change h_c d'un facteur 3 mais Q_tot de 1 % seulement, "
          "le film réfrigérant ne portant que 0,58 % de la résistance thermique totale (film gaz "
          "95,2 %, paroi 4,2 %). La température de paroi porte environ 54 K d'incertitude de "
          "fermeture."],
         ["Perte de charge côté calandre répartie uniformément sur les nœuds axiaux (branche monophasique)",
          "Bell-Delaware restitue une perte de charge globale de faisceau ; une répartition "
          "chicane par chicane exigerait de redériver ses termes par traversée",
          "Lorsque la température de saturation locale importe et que la perte de charge réelle est "
          "fortement non uniforme axialement"],
         ["Grandeurs de fermeture retardées d'un balayage (q″_w, Δp, T_wc)",
          "Brise une dépendance autrement circulaire entre les deux marches",
          "Ordre un dans le balayage ; s'annule à convergence mais contribue au cyclage limite "
          "lorsque la convergence stagne"],
         ["La pression du gaz n'est pas réinjectée dans la masse volumique du gaz",
          "La table FPV est construite à une pression de référence fixe",
          "Seulement si le nombre de Mach du gaz s'élève nettement au-dessus du plafond de "
          "dimensionnement de 0,3"],
         ["La branche monophasique côté calandre maintient la pression à sa valeur d'entrée",
          "L'enthalpie d'un gaz quasi parfait dépend faiblement de la pression",
          "La masse volumique et les propriétés de transport varient bien avec la pression ; la "
          "perte de charge Bell-Delaware calculée est actuellement écartée dans cette branche"],
         ["Le côté calandre est zéro-dimensionnel par nœud, sans équation de quantité de mouvement",
          "L'écoulement transversal chicané ne possède pas de coordonnée curviligne unique ; c'est "
          "la forme de modèle appropriée",
          "Ne constitue pas une limitation dans le régime modélisé"],
         ["Rayonnement du gaz désactivé",
          "Modèle implémenté mais désactivé dans ce chemin de solveur",
          "Température de gaz très élevée combinée à une forte émissivité"]],
        "Hypothèses du modèle et domaines de validité, classées approximativement par impact "
        "pratique sur les résultats restitués dans ce programme.",
        widths=[5.0, 5.6, 5.8])

    D.h2("11.1  Diagnostics auto-signalés")

    D.p("Plusieurs des limitations ci-dessus se manifestent à l'exécution plutôt que de rester "
        "implicites dans la documentation. À la lecture d'un calcul, les messages suivants sont "
        "porteurs d'information et doivent être consignés avec le résultat :")
    D.bullet("Avertissement d'extrapolation géométrique supercritique en écoulement transversal, "
             "avec le nom de la fermeture retenue.")
    D.bullet("Dépassement du critère de déclenchement de l'ébullition nucléée de Bergles-Rohsenow, "
             "avec la marge de surchauffe pariétale.")
    D.bullet("Avis de non-convergence du balayage, avec le résidu final.")
    D.bullet("NaN plutôt qu'une valeur fabriquée pour la marge de flux critique hors du dôme diphasique.")
    D.bullet("Régimes traversés et nombre de nœuds à risque de dégradation de l'échange thermique en "
             "fonctionnement supercritique.")
    D.bullet("Rapport de fuite aux chicanes r_lm de Bell-Delaware dépassant le domaine de calage de "
             "la corrélation.")
    D.bullet("Nombre de nœuds supercritiques dont le coefficient d'échange provient de Bell-Delaware "
             "plutôt que d'une fermeture à rapport de propriétés.")

    return D


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    fig_dir = os.environ.get("SHELLTUBE_FIG_DIR", here)
    out = os.path.join(here, "SHELLTUBE_STEADY_THEORY_FR.docx")
    first = build(fig_dir)                 # passe 1 : collecte des étiquettes
    doc = build(fig_dir, first.labels)     # passe 2 : résolution des renvois
    doc.save(out)
    print(f"écrit {out}")
    print(f"  équations : {doc.eq_n}   figures : {doc.fig_n}   tableaux : {doc.tab_n}")
    missing = [k for k, v in first.labels.items() if v is None]
    if missing:
        print("  ! étiquettes non résolues : " + ", ".join(missing), file=sys.stderr)


if __name__ == "__main__":
    main()
