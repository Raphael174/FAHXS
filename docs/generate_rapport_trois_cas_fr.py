"""Rapport français — comparaison thermodynamique de trois réfrigérants
sur l'échangeur calandre-et-tubes droits (régime stationnaire).

Réutilise la machinerie de mise en page de ``generate_shelltube_theory_docx``
(classe ``Doc``, convertisseur LaTeX -> OMML, tableaux, figures).

    SHELLTUBE_FIG_DIR=<dossier figures> python docs/generate_rapport_trois_cas_fr.py

Les données proviennent de ``trois_cas.json`` produit par le script d'étude ;
les figures sont attendues à 400 dpi dans le même dossier.
"""
from __future__ import annotations

import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from generate_shelltube_theory_docx import COLD, FLAG, HEAD_FONT, INK, SOFT, Doc  # noqa: E402
from docx.shared import Pt  # noqa: E402


def f(v, n=1):
    return ("%%.%df" % n % v).replace(".", ",")


def build(fig_dir, data):
    D = Doc()
    by = {d["label"]: d for d in data}
    he, n2, eau = by["Hélium"], by["LN2/N2"], by["Eau"]

    t = D.d.add_paragraph(); t.paragraph_format.space_after = Pt(2)
    r = t.add_run("Comportement thermodynamique — trois réfrigérants")
    r.font.name = HEAD_FONT; r.font.size = Pt(21); r.bold = True; r.font.color.rgb = INK

    s = D.d.add_paragraph(); s.paragraph_format.space_after = Pt(10)
    r = s.add_run("Échangeur calandre et tubes droits, Inconel 718 — régime stationnaire")
    r.font.name = HEAD_FONT; r.font.size = Pt(12.5); r.font.color.rgb = SOFT

    s = D.d.add_paragraph(); s.paragraph_format.space_after = Pt(12)
    r = s.add_run("Hélium · LN₂/N₂ supercritique · Eau avec changement de phase")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = COLD

    # ---------------- 1 ----------------
    D.h1("1.  Cas étudiés")

    D.p("Trois réfrigérants sont comparés sur la **même géométrie** (235 tubes droits "
        "Inconel 718, `shellTubeProp()` par défaut), en co-courant, chimie à vitesse finie "
        "(FPV), rapport de mélange OF = 2. Le débit de gaz chaud est imposé pour l'Hélium "
        "et ajusté pour les deux autres cas de façon à respecter la limite de validité "
        "Mach < 0,3 côté tubes.")

    D.table(
        ["Cas", "Réfrigérant", "T entrée", "p entrée", "Débit réfrigérant", "Débit gaz chaud"],
        [["1", "Hélium", "90 K", "85 bar", "150 g/s", "75 g/s (imposé)"],
         ["2", "LN₂ / N₂ supercritique", "90 K", "85 bar", "1 916 g/s", "65 g/s (ajusté Mach)"],
         ["3", "Eau", "300 K", "85 bar", "86 g/s", "65 g/s (ajusté Mach)"]],
        "Conditions d'entrée. Les trois cas emploient la fermeture fluide réel "
        "(`equilibrium_liquid`) : la sélection de corrélation côté calandre se fait ensuite "
        "sur l'état thermodynamique local, et non sur le fluide déclaré.",
        widths=[1.0, 4.0, 1.9, 1.9, 3.2, 4.2])

    D.note("Ajustement du débit de gaz chaud",
           "Le Mach côté tubes est maximal à l'entrée gaz (3020 K, gaz le moins dense) et "
           "décroît ensuite. À 65 g/s le maximum vaut 0,290 ; à 75 g/s il atteint 0,335, "
           "soit légèrement au-delà de la limite. Le cas Hélium, dont le débit est imposé "
           "à 75 g/s, dépasse donc marginalement le domaine de validité — voir §5.")

    # ---------------- 2 ----------------
    D.h1("2.  Synthèse des performances")

    D.table(
        ["Grandeur", "Hélium", "LN₂ / N₂", "Eau"],
        [["Puissance thermique Q", f(he["Q_kW"]) + " kW", f(n2["Q_kW"]) + " kW", f(eau["Q_kW"]) + " kW"],
         ["Température sortie réfrigérant", f(he["T_c_out"]) + " K", f(n2["T_c_out"]) + " K", f(eau["T_c_out"]) + " K"],
         ["Élévation ΔT réfrigérant", f(he["T_c_out"] - he["T_c_in"]) + " K",
          f(n2["T_c_out"] - n2["T_c_in"]) + " K", f(eau["T_c_out"] - eau["T_c_in"]) + " K"],
         ["Pression sortie", f(he["p_out_bar"], 2) + " bar", f(n2["p_out_bar"], 2) + " bar", f(eau["p_out_bar"], 2) + " bar"],
         ["Perte de charge côté calandre", f(he["dp_bar"], 2) + " bar", f(n2["dp_bar"], 2) + " bar", f(eau["dp_bar"], 2) + " bar"],
         ["dont terme d'accélération", f(he["dp_accel_bar"], 2) + " bar", f(n2["dp_accel_bar"], 2) + " bar", f(eau["dp_accel_bar"], 2) + " bar"],
         ["**Débit volumique en sortie**", "**" + f(he["Vdot_out_Ls"], 2) + " L/s**",
          "**" + f(n2["Vdot_out_Ls"], 2) + " L/s**", "**" + f(eau["Vdot_out_Ls"], 2) + " L/s**"],
         ["Masse volumique en sortie", f(he["rho_out"], 1) + " kg/m³", f(n2["rho_out"], 1) + " kg/m³", f(eau["rho_out"], 1) + " kg/m³"],
         ["**T paroi max (côté gaz)**", "**" + f(he["T_wg_max"]) + " K**", "**" + f(n2["T_wg_max"]) + " K**", "**" + f(eau["T_wg_max"]) + " K**"],
         ["T paroi max (côté réfrigérant)", f(he["T_wc_max"]) + " K", f(n2["T_wc_max"]) + " K", f(eau["T_wc_max"]) + " K"],
         ["État en sortie", "supercritique", "supercritique", "vapeur surchauffée, x = " + f(eau["quality_out"], 2)],
         ["Mach max côté gaz", f(he["Mach_max"], 3), f(n2["Mach_max"], 3), f(eau["Mach_max"], 3)],
         ["Marge flambage (ΔP/P_cr)", f(he["collapse"], 4), f(n2["collapse"], 4), f(eau["collapse"], 4)],
         ["Balayages jusqu'à convergence", str(he["n_sweeps"]), str(n2["n_sweeps"]), str(eau["n_sweeps"])]],
        "Synthèse. Rappel : la limite du domaine de données caractérisées de l'Inconel 718 "
        "est 1033 K, la tolérance admise pour cette étude 1500 K.",
        widths=[6.0, 3.4, 3.4, 3.6])

    D.figure(os.path.join(fig_dir, "fig1_comparaison.png"), 16.4,
             "Comparaison des trois cas. (a) températures de paroi, avec les maxima repérés — "
             "tous se situent au premier nœud, à l'entrée gaz. (b) échauffement du réfrigérant. "
             "(c) coefficients d'échange, en échelle logarithmique. (d) puissance et débit "
             "volumique délivré.")

    # ---------------- 3 ----------------
    D.h1("3.  Lecture physique")

    D.h2("3.1  Le film gaz contrôle l'échange")

    D.p("La figure 1(c) est la clé de lecture de toute l'étude : le coefficient d'échange "
        "côté réfrigérant dépasse celui du côté gaz d'un à deux ordres de grandeur, quel que "
        "soit le fluide. La résistance thermique est donc dominée par le film gaz "
        "(environ 77 à 95 % du total selon le cas), ce qui a deux conséquences directes :")
    D.bullet("Les trois fluides délivrent des puissances **voisines** — "
             + f(he["Q_kW"]) + ", " + f(n2["Q_kW"]) + " et " + f(eau["Q_kW"]) +
             " kW — malgré des débits massiques dans un rapport de 1 à 22. La puissance est "
             "fixée par ce que le gaz peut céder, pas par la capacité du réfrigérant à absorber.")
    D.bullet("Le choix de la corrélation d'ébullition ou supercritique influe peu sur la "
             "puissance : un écart de ±30 % sur le coefficient côté réfrigérant ne déplace "
             "le produit UA que de quelques pour cent.")

    D.h2("3.2  Ce qui distingue réellement les trois fluides")

    D.p("Ce n'est donc pas la puissance qui sépare les cas, mais **l'état de sortie** et la "
        "**température de paroi** :")

    D.table(
        ["Fluide", "Comportement", "Conséquence dimensionnante"],
        [["Hélium",
          "Monophasique supercritique sur toute la longueur. T_pc ≈ 11 K, très loin du "
          "domaine parcouru : le fluide se comporte comme un gaz parfait, c_p quasi constant.",
          "Débit volumique le plus élevé (" + f(he["Vdot_out_Ls"], 1) + " L/s) grâce à sa très "
          "faible masse volumique. Mais T paroi = " + f(he["T_wg_max"]) + " K, au-delà du domaine "
          "caractérisé, à cause d'un débit gaz plus élevé (75 g/s)."],
         ["LN₂ / N₂",
          "Supercritique de bout en bout (p_crit ≈ 34 bar). Traverse la région "
          "pseudo-critique : T_pc ≈ 148 K, la paroi la dépasse alors que le fluide de mélange "
          "reste en dessous.",
          "**Le plus favorable thermiquement** : T paroi = " + f(n2["T_wg_max"]) + " K, très en "
          "deçà de la limite matériau. Le fort débit (1 916 g/s) maintient la paroi froide. "
          "En contrepartie, la perte de charge est la plus élevée (" + f(n2["dp_bar"], 1) + " bar)."],
         ["Eau",
          "Traverse les trois régimes : liquide sous-refroidi, ébullition complète, puis "
          "surchauffe jusqu'à x = " + f(eau["quality_out"], 2) + ".",
          "Sortie vapeur surchauffée à " + f(eau["T_c_out"]) + " K, exploitable comme "
          "pressurant. Mais T paroi = " + f(eau["T_wg_max"]) + " K et **dépassement du flux "
          "critique** — voir §4."]],
        "Comportement et facteur dimensionnant par fluide.",
        widths=[2.0, 6.6, 7.8])

    # ---------------- 4 ----------------
    D.h1("4.  Cas Eau — changement de phase et flux critique")

    D.figure(os.path.join(fig_dir, "fig2_eau_changement_phase.png"), 14.0,
             "Cas Eau. (a) titre vapeur et découpage des régimes. (b) marge au flux critique, "
             "en échelle logarithmique. (c) température et effondrement de la masse volumique.")

    D.p("La colonne d'eau traverse intégralement le dôme diphasique :")
    D.bullet("**Liquide sous-refroidi** de 0 à ≈ 17 mm, le titre partant de x = −0,87.")
    D.bullet("**Ébullition** de ≈ 17 à 64,6 mm, où x atteint 1 (vaporisation complète).")
    D.bullet("**Vapeur surchauffée** de 64,6 mm à la sortie, x atteignant "
             + f(eau["quality_out"], 2) + " — la convention du code prolonge x au-delà de 1 pour "
             "quantifier la surchauffe, sans écrêtage à [0, 1].")

    D.p("Le palier de température visible en figure 2(c) entre 17 et 64 mm est la signature "
        "de l'ébullition : la température reste bloquée à T_sat(p) ≈ 573 K pendant que la "
        "chaleur latente est absorbée. C'est précisément ce que la marche en enthalpie permet "
        "de représenter et qu'une marche en température ne saurait restituer.")

    D.note("Dépassement du flux critique (CHF)",
           "La marge q″_CHF / q″_w passe **sous 1 entre 61,1 et 64,6 mm**, avec un minimum de "
           "0,013. Le titre y vaut déjà 0,953 : il s'agit donc de la transition d'assèchement "
           "(dryout) juste avant vaporisation complète, et non d'une crise d'ébullition "
           "prématurée en plein régime nucléé. Le flux massique côté calandre n'est que de "
           "296 kg/m²·s, ce qui abaisse mécaniquement le flux critique admissible. "
           "**Conséquence pratique** : sur ces quelques millimètres la paroi passe en régime "
           "post-CHF, où le modèle d'échange employé est une simplification conservative "
           "(vapeur saturée portant tout le débit, sans gouttelettes entraînées). La montée "
           "de température de paroi observée au-delà de 65 mm en est la traduction directe.",
           colour=FLAG)

    D.p("À noter : la perte de charge du cas Eau est la plus faible des trois "
        "(" + f(eau["dp_bar"], 2) + " bar) malgré la vaporisation, simplement parce que le débit "
        "massique est vingt fois inférieur à celui du cas N₂.")

    # ---------------- 5 ----------------
    D.h1("5.  Répartition de puissance et validité aérodynamique")

    D.figure(os.path.join(fig_dir, "fig3_flux_mach.png"), 16.4,
             "(a) répartition axiale de la puissance échangée. (b) nombre de Mach côté tubes, "
             "avec la limite de validité à 0,3.")

    D.p("La puissance est fortement concentrée à l'entrée gaz, où l'écart de température est "
        "maximal (3020 K contre un réfrigérant encore froid). C'est aussi là que se situe le "
        "maximum de température de paroi pour les trois cas.")

    D.note("Le cas Hélium sort du domaine de validité aérodynamique",
           "Le Mach côté tubes atteint " + f(he["Mach_max"], 3) + " pour l'Hélium, contre "
           + f(n2["Mach_max"], 3) + " pour les deux autres cas. Au-delà de 0,3 la marche de "
           "pression explicite du gaz (dp/dx = −f·ρ·U²/2D) fait chuter p_g rapidement, la masse "
           "volumique s'effondre et le calcul perd son sens ; c'est la raison d'être de cette "
           "limite. À 0,335 le dépassement reste marginal et le calcul demeure exploitable, "
           "mais un abaissement du débit gaz à 65 g/s le ramènerait dans le domaine, au prix "
           "d'environ 12 % de puissance.",
           colour=FLAG)

    # ---------------- 6 ----------------
    D.h1("6.  Réserves de modélisation")

    D.table(
        ["Réserve", "Portée"],
        [["Rapport de fuite aux chicanes r_lm = " + f(he["r_lm"], 2),
          "Le domaine de calage de Bell-Delaware est r_lm ≲ 1. L'aire de jeu tube-chicane vaut "
          "ici ≈ 6,5 fois l'aire de passage transversale (jeu de 1 mm sur un tube de 5 mm, "
          "chicanes espacées de 12 mm). Les deux corrections de fuite sont donc extrapolées. "
          "**C'est la première source d'incertitude sur la perte de charge**, devant tout choix "
          "de corrélation d'ébullition."],
         ["Corrélations d'ébullition en tube appliquées en écoulement transversal",
          "Gungor-Winterton est une corrélation d'écoulement interne. Une corrélation de "
          "faisceau serait plus défendable, mais l'effet est faible ici puisque le film "
          "réfrigérant ne porte qu'une fraction de la résistance totale (§3.1)."],
         ["Coefficient B de Chisholm reconstitué",
          "La table de B n'a pas pu être transcrite d'une source primaire. À vérifier avant de "
          "s'appuyer sur une perte de charge diphasique fine."],
         ["Températures de paroi au-delà de 1033 K",
          "Les cas Hélium (" + f(he["T_wg_max"]) + " K) et Eau (" + f(eau["T_wg_max"]) + " K) "
          "sortent du domaine de données caractérisées de l'Inconel 718. La courbe de limite "
          "d'élasticité y est écrêtée à plat : les marges mécaniques calculées au-delà ne sont "
          "pas étayées, même si elles restent sous la tolérance de 1500 K."]],
        "Réserves connues, par ordre d'importance pratique sur ces résultats.",
        widths=[5.2, 11.2])

    D.p("Les marges au flambage sous pression externe restent très confortables dans les trois "
        "cas (ΔP/P_cr de l'ordre de 0,002, pour un critère de sécurité à 1).")

    return D


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    fig_dir = os.environ.get("SHELLTUBE_FIG_DIR", here)
    data = json.load(open(os.path.join(fig_dir, "trois_cas.json"), encoding="utf-8"))
    out = os.path.join(here, "RAPPORT_TROIS_CAS_FR.docx")
    build(fig_dir, data).save(out)
    print("écrit " + out)


if __name__ == "__main__":
    main()
