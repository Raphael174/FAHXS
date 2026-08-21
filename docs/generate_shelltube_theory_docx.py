"""Generate the shell-and-tube steady solver theoretical-framework document (.docx).

Backbone reference for Helium / LN2 / Water heat-exchanger results. Regenerate
after any change to the solver's governing equations or closures:

    python docs/generate_shelltube_theory_docx.py

Equations are emitted as native Word (OMML) objects, not images, so they stay
editable in Word. A small LaTeX-subset parser (``_ml``) handles the constructs
used here: _{}, ^{}, \\frac, \\sqrt, \\sum, \\text, and greek/operator macros.

Figures are produced by the companion script kept alongside the generated
document; if the PNGs are absent the corresponding figure is skipped with a
warning rather than failing the build.
"""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --------------------------------------------------------------------------
# palette / metrics
# --------------------------------------------------------------------------
INK = RGBColor(0x1D, 0x21, 0x26)
SOFT = RGBColor(0x5C, 0x66, 0x72)
HOT = RGBColor(0xB4, 0x43, 0x2A)
COLD = RGBColor(0x14, 0x6C, 0x7E)
FLAG = RGBColor(0x9A, 0x5B, 0x12)

BODY_FONT = "Cambria"
HEAD_FONT = "Calibri Light"
MONO_FONT = "Consolas"

MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# --------------------------------------------------------------------------
# minimal LaTeX-subset -> OMML
# --------------------------------------------------------------------------
_GREEK = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ",
    "kappa": "κ", "lambda": "λ", "mu": "μ", "nu": "ν",
    "xi": "ξ", "pi": "π", "rho": "ρ", "sigma": "σ",
    "tau": "τ", "phi": "φ", "chi": "χ", "psi": "ψ",
    "omega": "ω", "Delta": "Δ", "Sigma": "Σ", "Phi": "Φ",
    "Omega": "Ω", "Gamma": "Γ", "Theta": "Θ",
    "varphi": "φ", "varepsilon": "ε", "vartheta": "ϑ",
}
_OPS = {
    "cdot": "·", "times": "×", "pm": "±", "approx": "≈",
    "leq": "≤", "geq": "≥", "neq": "≠", "to": "→",
    "leftarrow": "←", "Rightarrow": "⇒", "infty": "∞",
    "partial": "∂", "propto": "∝", "ll": "≪", "gg": "≫",
    "quad": " ", "qquad": "  ", "," : " ", ";": " ",
    "equiv": "≡", "sim": "∼", "in": "∈",
}
# Upright multi-letter function names (operators are roman in math typography).
_FUNCS = {"ln", "log", "exp", "max", "min", "sin", "cos", "tan",
          "sinh", "cosh", "tanh", "lim", "det"}
# \dot{} etc -> OMML accent (combining mark above); \bar{} -> OMML overbar
_ACCENTS = {"dot": "̇", "ddot": "̈", "hat": "̂",
            "tilde": "̃", "vec": "⃗"}
_LETTER = re.compile(r"[A-Za-zͰ-Ͽ]")


def _esc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _run(text: str, upright: bool) -> str:
    pr = "<m:rPr><m:nor/></m:rPr>" if upright else ""
    return f'<m:r>{pr}<m:t xml:space="preserve">{_esc(text)}</m:t></m:r>'


def _text_atoms(t: str) -> str:
    """Letters italic (variables), digits/operators upright — correct math style."""
    if not t:
        return ""
    parts = re.findall(r"[A-Za-zͰ-Ͽ]+|[^A-Za-zͰ-Ͽ]+", t)
    return "".join(_run(p, upright=not _LETTER.match(p)) for p in parts)


def _group(s: str, i: int):
    if i < len(s) and s[i] == "{":
        depth, j = 1, i + 1
        while j < len(s) and depth:
            if s[j] == "{":
                depth += 1
            elif s[j] == "}":
                depth -= 1
            j += 1
        return s[i + 1:j - 1], j
    return (s[i] if i < len(s) else ""), i + 1


def _split_right(s: str, i: int):
    """Consume '<delim> ... \\right<delim>' starting just after '\\left'.

    Returns (inner, begChr, endChr, index_after_right). Handles nesting.
    """
    beg = s[i] if i < len(s) else "("
    i += 1
    depth, j = 1, i
    while j < len(s):
        m = re.match(r"\\(left|right)", s[j:])
        if m:
            if m.group(1) == "left":
                depth += 1
                j += 5
            else:
                depth -= 1
                if depth == 0:
                    inner = s[i:j]
                    j += 6                       # past '\right'
                    end = s[j] if j < len(s) else ")"
                    return inner, beg, end, j + 1
                j += 6
            continue
        j += 1
    return s[i:], beg, ")", len(s)               # unterminated: degrade gracefully


def _ml(s: str) -> str:
    """Convert a LaTeX-subset string to an OMML fragment."""
    out, buf, i = [], [], 0

    def flush():
        if buf:
            out.append(_text_atoms("".join(buf)))
            buf.clear()

    while i < len(s):
        c = s[i]
        if c == "\\":
            m = re.match(r"\\([a-zA-Z]+|,|;)", s[i:])
            if not m:
                i += 1
                continue
            cmd = m.group(1)
            i += 1 + len(cmd)
            if cmd in _GREEK:
                buf.append(_GREEK[cmd]); continue
            if cmd in _OPS:
                buf.append(_OPS[cmd]); continue
            if cmd in _FUNCS:
                flush()
                out.append(_run(cmd, upright=True))
                continue
            if cmd in _ACCENTS:
                flush()
                a, i = _group(s, i)
                out.append(f'<m:acc><m:accPr><m:chr m:val="{_ACCENTS[cmd]}"/></m:accPr>'
                           f"<m:e>{_ml(a)}</m:e></m:acc>")
                continue
            if cmd == "bar":
                flush()
                a, i = _group(s, i)
                out.append('<m:bar><m:barPr><m:pos m:val="top"/></m:barPr>'
                           f"<m:e>{_ml(a)}</m:e></m:bar>")
                continue
            if cmd == "left":
                flush()
                inner, beg, end, i = _split_right(s, i)
                out.append(f'<m:d><m:dPr><m:begChr m:val="{beg}"/>'
                           f'<m:endChr m:val="{end}"/></m:dPr>'
                           f"<m:e>{_ml(inner)}</m:e></m:d>")
                continue
            flush()
            if cmd == "frac":
                a, i = _group(s, i); b, i = _group(s, i)
                out.append(f"<m:f><m:num>{_ml(a)}</m:num><m:den>{_ml(b)}</m:den></m:f>")
            elif cmd == "sqrt":
                a, i = _group(s, i)
                out.append("<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr>"
                           f"<m:deg/><m:e>{_ml(a)}</m:e></m:rad>")
            elif cmd == "text":
                a, i = _group(s, i)
                out.append(_run(a, upright=True))
            elif cmd == "sum":
                sb = sp = ""
                if i < len(s) and s[i] == "_":
                    i += 1; a, i = _group(s, i); sb = _ml(a)
                if i < len(s) and s[i] == "^":
                    i += 1; a, i = _group(s, i); sp = _ml(a)
                a, i = _group(s, i)
                hide = "1" if not sp else "0"
                out.append(
                    f'<m:nary><m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/>'
                    f'<m:supHide m:val="{hide}"/></m:naryPr><m:sub>{sb}</m:sub>'
                    f'<m:sup>{sp}</m:sup><m:e>{_ml(a)}</m:e></m:nary>')
            elif cmd in ("left", "right"):
                pass
            continue

        if c in "_^":
            if buf:
                base_ch = buf.pop()
                flush()
                base = _text_atoms(base_ch)
            elif out:
                base = out.pop()
            else:
                base = ""
            i += 1
            a, i = _group(s, i)
            inner = _ml(a)
            if c == "_" and i < len(s) and s[i] == "^":
                i += 1
                b, i = _group(s, i)
                out.append(f"<m:sSubSup><m:e>{base}</m:e><m:sub>{inner}</m:sub>"
                           f"<m:sup>{_ml(b)}</m:sup></m:sSubSup>")
            elif c == "_":
                out.append(f"<m:sSub><m:e>{base}</m:e><m:sub>{inner}</m:sub></m:sSub>")
            else:
                out.append(f"<m:sSup><m:e>{base}</m:e><m:sup>{inner}</m:sup></m:sSup>")
            continue

        buf.append(c)
        i += 1

    flush()
    return "".join(out)


# --------------------------------------------------------------------------
# document helpers
# --------------------------------------------------------------------------
class Doc:
    def __init__(self, refmap=None):
        self.d = Document()
        self.eq_n = 0
        self.fig_n = 0
        self.tab_n = 0
        self.labels = {}                 # label -> equation number (this pass)
        self.refmap = refmap or {}       # label -> number (previous pass)
        self._setup()

    def ref(self, label):
        """Resolve an equation cross-reference. Two-pass: empty on pass 1."""
        n = self.refmap.get(label)
        if n is None:
            self.labels.setdefault(label, None)
            return "Eq. --"
        return f"Eq. {n}"

    def _setup(self):
        sec = self.d.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
        sec.left_margin = sec.right_margin = Cm(2.3)
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)

        st = self.d.styles["Normal"]
        st.font.name = BODY_FONT
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.12

        for name, size, colour in (("Heading 1", 16, INK),
                                   ("Heading 2", 12.5, INK),
                                   ("Heading 3", 11, SOFT)):
            s = self.d.styles[name]
            s.font.name = HEAD_FONT
            s.font.size = Pt(size)
            s.font.color.rgb = colour
            s.font.bold = True
            s.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
            s.paragraph_format.space_after = Pt(5)
            s.paragraph_format.keep_with_next = True

    # -- text ------------------------------------------------------------
    def h1(self, t): return self.d.add_heading(t, 1)
    def h2(self, t): return self.d.add_heading(t, 2)
    def h3(self, t): return self.d.add_heading(t, 3)

    def p(self, text="", *, italic=False, size=None, colour=None, space_after=None):
        par = self.d.add_paragraph()
        if text:
            self._rich(par, text, italic=italic, size=size, colour=colour)
        if space_after is not None:
            par.paragraph_format.space_after = Pt(space_after)
        return par

    def _rich(self, par, text, *, italic=False, size=None, colour=None):
        """**bold**, `mono` and plain segments."""
        for seg in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
            if not seg:
                continue
            if seg.startswith("**") and seg.endswith("**"):
                r = par.add_run(seg[2:-2]); r.bold = True
            elif seg.startswith("`") and seg.endswith("`"):
                r = par.add_run(seg[1:-1]); r.font.name = MONO_FONT
                r.font.size = Pt((size or 10.5) - 1.0)
            else:
                r = par.add_run(seg)
            r.italic = italic
            if size:
                r.font.size = Pt(size)
            if colour:
                r.font.color.rgb = colour
        return par

    def bullet(self, text, level=0):
        par = self.d.add_paragraph(style="List Bullet")
        par.paragraph_format.left_indent = Cm(0.7 + 0.6 * level)
        par.paragraph_format.space_after = Pt(3)
        self._rich(par, text)
        return par

    def note(self, title, text, colour=COLD):
        par = self.d.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.5)
        par.paragraph_format.space_before = Pt(8)
        par.paragraph_format.space_after = Pt(8)
        r = par.add_run(title.upper() + "  ")
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = colour
        r.font.name = HEAD_FONT
        self._rich(par, text, size=9.5)
        pPr = par._p.get_or_add_pPr()
        pPr.append(parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:left w:val="single" w:sz="18" w:space="10" w:color="{_hex(colour)}"/></w:pBdr>'))
        return par

    # -- equations -------------------------------------------------------
    def eq(self, latex, *, number=True, tag=None, label=None):
        if number:
            self.eq_n += 1
            if label:
                self.labels[label] = self.eq_n
        par = self.d.add_paragraph()
        pf = par.paragraph_format
        pf.space_before = Pt(7); pf.space_after = Pt(7)
        pf.tab_stops.add_tab_stop(Cm(7.7), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(16.4), WD_TAB_ALIGNMENT.RIGHT)
        par.add_run("\t")
        par._p.append(parse_xml(f'<m:oMath xmlns:m="{MNS}">{_ml(latex)}</m:oMath>'))
        if number:
            r = par.add_run("\t(" + (tag or str(self.eq_n)) + ")")
            r.font.size = Pt(9.5); r.font.color.rgb = SOFT
        return par

    def eqs(self, *latex_lines, tag=None, label=None):
        """A numbered group: several stacked lines, number on the first."""
        for k, line in enumerate(latex_lines):
            self.eq(line, number=(k == 0), tag=tag, label=label if k == 0 else None)

    def caption(self, kind, text):
        if kind == "fig":
            self.fig_n += 1; n = self.fig_n
        else:
            self.tab_n += 1; n = self.tab_n
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_after = Pt(11)
        r = par.add_run(f"{'Figure' if kind == 'fig' else 'Table'} {n}.  ")
        r.bold = True; r.font.size = Pt(9)
        self._rich(par, text, size=9, colour=SOFT)
        return par

    def figure(self, png, width_cm, caption):
        if not os.path.exists(png):
            print(f"  ! missing figure {png} — skipped", file=sys.stderr)
            return
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(8)
        par.paragraph_format.space_after = Pt(4)
        par.add_run().add_picture(png, width=Cm(width_cm))
        self.caption("fig", caption)

    # -- tables ----------------------------------------------------------
    def table(self, headers, rows, caption, widths=None, mono_cols=()):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, h in enumerate(headers):
            hdr[j].text = ""
            par = hdr[j].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run(h)
            r.bold = True; r.font.size = Pt(8.5); r.font.name = HEAD_FONT
            _shade(hdr[j], "EDEFF2")
        for row in rows:
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = ""
                par = cells[j].paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                if j in mono_cols:
                    r = par.add_run(v); r.font.name = MONO_FONT; r.font.size = Pt(8.5)
                else:
                    self._rich(par, v, size=9)
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Cm(w)
        self.caption("tab", caption)
        return t

    def save(self, path):
        self.d.save(path)


def _hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def _shade(cell, hex_fill):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_fill}"/>'))


# ==========================================================================
# document content
# ==========================================================================
def build(fig_dir: str, refmap=None) -> Doc:
    D = Doc(refmap)

    # ---------------- title block ----------------
    t = D.d.add_paragraph()
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("Shell-and-Tube Steady Solver")
    r.font.name = HEAD_FONT; r.font.size = Pt(24); r.bold = True; r.font.color.rgb = INK

    s = D.d.add_paragraph()
    s.paragraph_format.space_after = Pt(10)
    r = s.add_run("Theoretical Framework, Governing Equations and Empirical Closures")
    r.font.name = HEAD_FONT; r.font.size = Pt(13); r.font.color.rgb = SOFT

    s = D.d.add_paragraph()
    s.paragraph_format.space_after = Pt(14)
    r = s.add_run("Backbone reference for Helium / LN₂ / Water heat-exchanger results")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = COLD

    D.table(
        ["Item", "Value"],
        [["Model", "Quasi-1D steady baffled shell-and-tube combustor heat exchanger"],
         ["Configuration", "Hot combustion gas inside straight tubes; coolant in shell-side cross-flow"],
         ["Primary implementation", "1Dmodel/main_solve_shellntube.py"],
         ["Wall model", "physics/heat_conduction.py — OneDimensionalSteadyConduction_ShellnHelicalTube"],
         ["Shell-side closure", "physics/bell_delaware.py (single phase); physics/liquid_flow/ (two-phase, supercritical)"],
         ["Entry point", "1Dmodel/main_steady.py — combustorProp.HX_config = \"shellntube\""],
         ["Units", "SI throughout, except material property tables which are in °C"]],
        "Document scope and source of record. Regenerate this document with "
        "`docs/generate_shelltube_theory_docx.py` after any change to the governing equations or closures.",
        widths=[4.2, 12.0])

    # ---------------- 1 ----------------
    D.h1("1.  Scope and physical configuration")

    D.p("Hot combustion products from a liquid-fuel/oxygen combustor flow axially inside a bundle "
        "of straight tubes. Coolant flows on the shell side, forced into a zig-zag cross-flow path "
        "by segmental baffles. Heat passes radially outward from the gas, through the tube wall, "
        "into the coolant.")

    D.p("This document specifies the model as implemented, including the empirical correlations and "
        "the assumptions attached to each. It is written to be the common technical basis for the "
        "Helium, LN₂/supercritical-N₂ and Water results, which differ only in which coolant-side "
        "branch of the model is exercised — the gas-side, wall and iteration structure are identical "
        "for all three. Fluid-specific consequences are collected in Section 10.")

    D.h2("1.1  Nomenclature")
    D.table(
        ["Symbol", "Quantity", "Units"],
        [["A_i, A_hot, A_cold", "Tube inner flow area; nodal hot- and cold-face areas", "m²"],
         ["B", "Baffle spacing", "m"],
         ["Bo", "Boiling number, q″/(G h_fg)", "–"],
         ["cp", "Specific heat at constant pressure", "J/(kg·K)"],
         ["D_i, D_o", "Tube inner and outer diameter", "m"],
         ["D_s, D_otl", "Shell inner diameter; outer tube limit diameter", "m"],
         ["dQ_i", "Heat rate at node i, per single tube", "W"],
         ["f", "Darcy friction factor", "–"],
         ["G_s, G_w", "Shell-side cross-flow and window mass flux", "kg/(m²·s)"],
         ["h", "Specific enthalpy", "J/kg"],
         ["h_g, h_c", "Gas- and coolant-side film coefficients", "W/(m²·K)"],
         ["h_fg", "Latent heat of vaporization", "J/kg"],
         ["j", "Colburn j-factor", "–"],
         ["k_w", "Wall thermal conductivity", "W/(m·K)"],
         ["L_tube", "Tube length", "m"],
         ["ṁ_g, ṁ_c, ṁ_tube", "Hot-gas, coolant and per-tube gas mass flow", "kg/s"],
         ["N, N_tubes, N_b", "Axial nodes; tubes in bundle; baffles", "–"],
         ["N_tcc, N_tcw", "Tube rows crossed in one cross-flow section; in one window", "–"],
         ["Nu, Pr, Re", "Nusselt, Prandtl, Reynolds numbers", "–"],
         ["p, p_r", "Pressure; reduced pressure p/p_crit", "Pa, –"],
         ["P_t, P_n", "Tube pitch; pitch normal to flow", "m"],
         ["q″_w", "Wall heat flux, outer-area basis", "W/m²"],
         ["R_wall", "Nodal cylindrical wall thermal resistance", "K/W"],
         ["S_m, S_w, S_tb, S_sb, S_b", "Bell-Delaware cross-flow, window, tube-baffle, shell-baffle and bypass areas", "m²"],
         ["s_w", "Tube wall thickness", "m"],
         ["T_g, T_c", "Bulk gas and coolant temperature", "K"],
         ["T_wg, T_wc", "Gas-side and coolant-side wall face temperature", "K"],
         ["UA", "Nodal overall conductance", "W/K"],
         ["U_g", "Gas velocity", "m/s"],
         ["x", "Thermodynamic quality (vapour mass fraction)", "–"],
         ["X_tt", "Lockhart-Martinelli parameter, turbulent-turbulent", "–"],
         ["α", "Void fraction", "–"],
         ["Δx", "Axial cell length", "m"],
         ["μ, ρ, ν", "Viscosity, density, Poisson ratio", "Pa·s, kg/m³, –"],
         ["φ", "Corrugation severity index, e²/(p D_i)", "–"],
         ["ω", "Under-relaxation factor", "–"]],
        "Principal symbols. Subscripts g and c denote the hot gas and the coolant respectively; "
        "l and v denote saturated liquid and vapour; b and w denote bulk and wall conditions.",
        widths=[4.4, 9.4, 2.4], mono_cols=(0,))

    # ---------------- 2 ----------------
    D.h1("2.  Structure of the problem")

    D.p("The boundary conditions determine both the mathematical character of the problem and the "
        "solution strategy. Both fluid inlets are prescribed: the gas enters at the tube inlet with a "
        "temperature and composition fixed by the upstream combustion solve, and the coolant enters "
        "at its own inlet with a prescribed temperature and pressure. Crucially, **where** those two "
        "inlets sit relative to one another depends on the flow configuration.")

    D.table(
        ["Configuration", "Gas inlet", "Coolant inlet", "Mathematical character"],
        [["Co-flow", "node 0", "node 0",
          "**Initial-value problem.** All data is available at one end; a single simultaneous forward "
          "march would suffice in principle."],
         ["Counter-flow", "node 0", "node N−1",
          "**Two-point boundary value problem.** At node 0 the gas temperature is known but the coolant "
          "temperature is not — that is the coolant outlet, the one quantity not prescribed."]],
        "Boundary-condition structure by flow configuration. Counter-flow is what makes an iterative "
        "solution necessary; the same solution path is used for co-flow rather than maintaining a "
        "second code route.",
        widths=[2.9, 2.0, 2.2, 9.1])

    D.p("A two-point boundary value problem cannot be advanced by a single forward march, because half "
        "the required data lies at the far end of the domain. The solver therefore decomposes the "
        "problem into two one-way marches — a tube-side march and a shell-side march — and "
        "iterates them against one another until they are mutually consistent. This is the "
        "**predictive sweep** of Section 8.")

    D.note("Why the decomposition is valid",
           "Each march is a well-posed initial-value problem in isolation. Given a frozen shell-side "
           "temperature field, the tube-side march is a pure forward integration; given a frozen duty "
           "field, the shell-side energy march likewise. Only the coupling between them is implicit, "
           "and it is precisely that coupling which the outer iteration resolves.")

    D.p("A consequence worth recording: because the two marches are already decoupled, changing flow "
        "direction is inexpensive. Counter-flow requires only that the shell-side march loop be "
        "reversed, together with correct tracking of the coolant's own entrance distance for "
        "entrance-length corrections. No shooting method or root-find is required. This differs from "
        "the helical-coil solver in the same codebase, which marches wall and coolant together in one "
        "coupled pass and must therefore resolve counter-flow by shooting on the hot-end starting "
        "enthalpy with adaptive bracketing and bisection.")

    # ---------------- 3 ----------------
    D.h1("3.  Discretization")

    D.p("A uniform axial grid of N cells spans the tube length, with node i centred at "
        "x = (i + ½)Δx:")
    D.eq(r"\Delta x = \frac{L_{tube}}{N} \qquad N = 200 \text{ (default)}")

    D.p("Node index 0 always denotes the **gas-inlet end**, irrespective of coolant direction. "
        "Co-flow means the coolant also begins at index 0; counter-flow means it begins at index "
        "N−1 and marches downward. Maintaining a single index convention for both fluids allows "
        "one duty array to serve both marches without reindexing.")

    D.p("Tube-side geometry follows from the outer diameter and wall thickness, and the gas mass flow "
        "is divided across the bundle:")
    D.eqs(r"D_i = D_o - 2 s_w",
          r"A_i = \frac{\pi D_i^{2}}{4}",
          r"\dot{m}_{tube} = \frac{\dot{m}_g}{N_{tubes}}")

    D.note("Per-tube versus bundle quantities",
           "Every tube-side quantity in this model — velocity, Reynolds number, enthalpy budget and "
           "the nodal duty dQ_i — is defined **per single representative tube**. Multiplication by "
           "N_tubes occurs exactly once, where the duty crosses to the shell side (" + D.ref("ntubes") + "). Using the "
           "total gas mass flow in the tube-side velocity is a common and silent error.",
           colour=FLAG)

    D.note("S_sb corrected 2026-08-20",
           "The shell-to-baffle leakage area is the diametral gap over the arc where the "
           "baffle edge follows the shell: S_sb = Lsb*(Ds/2)*(2*pi - theta_ds) = "
           "Ds*Lsb*(pi - theta_ds/2). The implementation previously carried a spurious extra "
           "factor of 1/2. Confirmed against the first-principles gap-times-arc derivation "
           "and against Hellborg (2017) eq. 47, which agree with each other. Effect on this "
           "geometry: S_sb 1.218 -> 2.436 cm2, r_lm 6.51 -> 6.93.",
           colour=FLAG)

    D.p("On the shell side the geometry is reduced to the Bell-Delaware area set. The governing area "
        "is the minimum cross-flow area at the shell centreline, which sets the reference mass flux "
        "for every shell-side correlation:")
    D.eq(r"S_m = B \left[ (D_s - D_{otl}) + \frac{D_{otl} - D_o}{P_n}(P_t - D_o) \right]")

    # ---------------- 4 ----------------
    D.h1("4.  Tube-side (hot gas) model")

    D.p("A forward march from node 0 to node N−1, evaluated against the current shell-side "
        "temperature field.")

    D.h2("4.1  Enthalpy-removal state parameterization")

    D.p("Rather than carrying temperature as the state variable, the march carries the **cumulative "
        "specific enthalpy removed** from the gas. In the default finite-rate mode the complete "
        "thermochemical state is retrieved from a precomputed flamelet/progress-variable (FPV) "
        "manifold indexed by two scalars — the removed enthalpy and an unnormalized recombination "
        "progress variable Y_c:")
    D.eq(r"(T_g, \rho_g, \mu_g, k_g, c_{p,g}, \omega_{Yc}) = M(h_{removed}, Y_c)")

    D.p("Interpolation on the manifold is bilinear. The progress variable is transported by its own "
        "net production rate, which is what allows the composition to lag equilibrium as the gas "
        "cools; that lag is the finite-rate physics:")
    D.eqs(r"\frac{dY_c}{dx} = \frac{\omega_{Yc}(h_{removed}, Y_c)}{U_g}",
          r"\frac{dh_{removed}}{dx} = \frac{1}{\dot{m}_{tube}} \frac{dQ}{dx}")

    D.p("Both are advanced explicitly, node by node. Two comparison modes replace this manifold "
        "lookup with a direct chemical-equilibrium solver call at each node, re-equilibrating the "
        "mixture at constant pressure after the enthalpy removal (equilibrium mode) or holding "
        "composition fixed (frozen mode). Finite-rate is the physically appropriate default for the "
        "high heat-extraction regime of interest; the other two are validation modes only.")

    D.h2("4.2  Dimensionless groups and the gas-side film")

    D.eqs(r"U_g = \frac{\dot{m}_{tube}}{\rho_g A_i}",
          r"Re_g = \frac{\rho_g U_g D_i}{\mu_g}",
          r"Pr_g = \frac{c_{p,g}\, \mu_g}{k_g}",
          r"h_g = \frac{Nu_g\, k_g}{D_i}")

    D.p("The Nusselt number and friction factor are dispatched on the tube surface treatment. For "
        "grooved (corrugated) tubes, both are taken from the Vicente–Cruz corrugated-tube "
        "correlations, parameterized by a corrugation severity index built from groove depth e and "
        "groove pitch p:")
    D.eq(r"\varphi = \frac{e^{2}}{p\, D_i}")

    D.p("Smooth tubes instead use the standard straight-tube Nusselt dispatch with Colebrook "
        "friction. In both cases the result is scaled by calibration multipliers held in the "
        "`CorrelationCoefficients` dataclass, which is the single source of tuning knobs for the model.")

    D.h2("4.3  Gas-side momentum")

    D.eq(r"\frac{dp_g}{dx} = - \frac{f_g \rho_g U_g^{2}}{2 D_i}")

    D.note("Friction factor convention",
           "f_g is the **Darcy** friction factor throughout the maintained solver paths, consistent "
           "with the factor of two in the denominator above. Applying a Fanning-to-Darcy conversion "
           "on top of this quadruples the predicted pressure drop.",
           colour=FLAG)

    D.p("The coupling here is deliberately one-way: gas pressure is integrated for reporting but is "
        "not fed back into the gas density, because the FPV manifold is constructed at a "
        "fixed reference pressure. This is admissible because the combustor is designed for gas Mach "
        "numbers below approximately 0.3, where the density error incurred by neglecting the pressure "
        "drop is small. It is also the reason no quasi-one-dimensional momentum equation is solved "
        "for the hot gas anywhere in this model.")

    # ---------------- 5 ----------------
    D.h1("5.  One-dimensional wall heat-transfer model")

    D.p("At each node the two convective films and the tube wall form a three-resistance series "
        "network between the gas and coolant bulk temperatures. The essential geometric feature is "
        "that the two films act on **different areas**, because the tube is cylindrical and the hot "
        "fluid is on the inside.")

    D.figure(os.path.join(fig_dir, "fig_wall_network.png"), 16.2,
             "Radial arrangement and the equivalent nodal resistance network. Because the hot fluid is "
             "inside the tube, the gas-side film acts on the smaller inner perimeter and the coolant-side "
             "film on the larger outer perimeter.")

    D.h2("5.1  Geometry and the hot-side mapping")

    D.eqs(r"P_{hot} = \pi D_i \qquad P_{cold} = \pi D_o",
          r"A_{hot} = P_{hot}\, \Delta x \qquad A_{cold} = P_{cold}\, \Delta x")

    D.note("Hot-side orientation",
           "The wall model is shared with the helical-coil configuration, in which the hot fluid is "
           "**outside** the tube and the perimeter assignment is reversed. The orientation is selected "
           "by a `hot_side` flag, which must be set to \"inner\" for shell-and-tube. Setting it "
           "incorrectly leaves the network structurally intact but scales both film resistances by "
           "D_o/D_i in the wrong direction, producing a plausible but wrong duty.",
           colour=FLAG)

    D.h2("5.2  Wall resistance and overall conductance")

    D.p("The wall resistance is **cylindrical**, not planar — an important distinction at the wall "
        "thickness-to-diameter ratios used here:")
    D.eq(r"R_{wall} = \frac{\ln\left[(D_i/2 + s_w)/(D_i/2)\right]}{2 \pi \Delta x\, k_w}")

    D.p("The three resistances combine in series to give the nodal conductance and the nodal heat "
        "rate:")
    D.eqs(r"\frac{1}{UA} = \frac{1}{h_g A_{hot}} + R_{wall} + \frac{1}{h_c A_{cold}}",
          r"dQ_i = UA \left( T_{g,i} - T_{c,i} \right)")

    D.h2("5.3  Face temperatures")

    D.p("The wall face temperatures follow by walking the resistance chain in order, with the third "
        "relation acting as a consistency check that the chain returns to the prescribed coolant "
        "temperature:")
    D.eqs(r"T_{wg} = T_g - \frac{dQ}{h_g A_{hot}}",
          r"T_{wc} = T_{wg} - dQ\, R_{wall}",
          r"T_c^{check} = T_{wc} - \frac{dQ}{h_c A_{cold}} \;\;\equiv\;\; T_c")

    D.h2("5.4  Temperature-dependent conductivity and the nodal solve")

    D.p("Wall conductivity is evaluated at the through-thickness mean temperature:")
    D.eq(r"k_w = k_w(\bar{T}_w), \qquad \bar{T}_w = \frac{T_{wg} + T_{wc}}{2}")

    D.p("The wall resistance therefore depends on the very face temperatures being solved for. Each "
        "node consequently closes a small nonlinear system, solved to a tolerance of 10⁻⁸ on "
        "the residual vector:")
    D.eq(r"F = \left[ T_{wg} - T_{wg}^{new},\; T_{wc} - T_{wc}^{new},\; T_c^{check} - T_c^{check,new} \right] = 0")

    D.p("This nodal nonlinear solve is nested inside the axial node loop, which is in turn nested "
        "inside the outer sweep of Section 8 — three levels of iteration in total.")

    D.h2("5.5  Wall heat flux and radiation")

    D.p("The wall heat flux passed to the shell-side closures is referenced to the **outer** tube "
        "area, since that is the surface the coolant sees:")
    D.eq(r"q''_{w,i} = \frac{dQ_i}{\pi D_o \Delta x}")

    D.p("A gas radiation model (mean-beam-length gas emissivity, entering as a radiative coefficient "
        "in parallel with the convective gas-side film) is implemented in the wall module but is "
        "**disabled** in the shell-and-tube steady path.")

    # ---------------- 6 ----------------
    D.h1("6.  Shell-side heat transfer and pressure drop")

    D.h2("6.1  Model altitude")

    D.p("Shell-side flow crosses the tube bundle transversely, turns through a baffle window, and "
        "crosses again, alternating along the exchanger. Over most of its path the velocity vector is "
        "perpendicular to the axial coordinate, and there is no single streamwise coordinate along "
        "which a one-dimensional momentum equation could be posed. The shell side is therefore "
        "modelled as **zero-dimensional per axial node**: a lumped bundle correlation evaluated with "
        "local fluid properties. This is not a simplification of a momentum ODE — it is a "
        "different and more appropriate model form, empirically calibrated on precisely this "
        "geometry.")

    D.h2("6.2  Bell-Delaware ideal tube bank")

    D.eqs(r"G_s = \frac{\dot{m}_c}{S_m}",
          r"Re_s = \frac{D_o G_s}{\mu_s}",
          r"Pr_s = \frac{c_p \mu}{k}")

    D.p("The ideal-bank Colburn j-factor uses layout- and Reynolds-dependent piecewise coefficients "
        "a₁–a₄ tabulated for triangular, square and rotated-square layouts:")
    D.eqs(r"a = \frac{a_3}{1 + 0.14\, Re_s^{a_4}}",
          r"j = a_1 \left( \frac{1.33}{P_t/D_o} \right)^{a} Re_s^{a_2}",
          r"h_{ideal} = j\, c_p\, G_s\, Pr_s^{-2/3} \left( \frac{\mu_b}{\mu_w} \right)^{0.14}")

    D.note("The property correction is evaluated, not defaulted",
           "The (mu_b/mu_w)^0.14 Sieder-Tate term above is evaluated from the lagged coolant-side "
           "wall temperature (Section 8.1). It was previously left at its neutral default of 1.0, "
           "i.e. the correlation's own property-variation correction was switched off. Enabling it "
           "lowered the water design point's peak wall temperature by 67 K and resolved a "
           "convergence stall at that point; it moves the Helium shell-and-tube baseline by "
           "about 0.1 %. The ratio is clamped to [0.25, 4] so an unconverged cold-start wall "
           "state cannot distort the coefficient.")

    D.p("Five multiplicative correction factors then account for the real bundle's departures from an "
        "ideal bank:")
    D.eq(r"h_{shell} = h_{ideal}\, J_c\, J_l\, J_b\, J_s\, J_r", label="hshell")

    D.table(
        ["Factor", "Physical effect", "Form"],
        [["J_c", "Baffle-window configuration: tubes within the window experience axial rather than "
                 "cross flow. F_c is the fraction of tubes in pure cross-flow.",
          "0.55 + 0.72 F_c"],
         ["J_l", "Baffle leakage through tube-to-baffle and shell-to-baffle clearances, which bypasses "
                 "the bundle. r_lm = (S_sb+S_tb)/S_m, r_s = S_sb/(S_sb+S_tb).",
          "p + (1−p)exp(−2.2 r_lm),  p = 0.44(1−r_s)"],
         ["J_b", "Bundle bypass around the outside of the tube field. C = 1.35 laminar, 1.25 turbulent; "
                 "unity when the sealing-strip ratio r_ss ≥ 0.5.",
          "exp(−C F_sbp [1−(2 r_ss)^(1/3)])"],
         ["J_s", "Unequal inlet and outlet end baffle spacing relative to central spacing.",
          "spacing ratio, exponent n₁ = 0.6"],
         ["J_r", "Laminar adverse temperature gradient. Active only below Re ≈ 100, unity above.",
          "(10/N_c)^0.18, blended linearly to 1 at Re = 100"]],
        "Bell-Delaware correction factors. Typical magnitudes for a 25 % baffle cut are J_c ≈ 1.0, "
        "J_l ≈ 0.7–0.8, J_b ≈ 0.9, J_s ≈ 1 and J_r = 1 in turbulent flow.",
        widths=[1.6, 9.0, 5.8], mono_cols=(0,))

    D.h2("6.3  Shell-side pressure drop")

    D.p("Pressure drop is assembled from three zones — the central cross-flow passes, the baffle "
        "windows, and the two end zones — each carrying its own leakage and bypass correction "
        "(R_l, R_b, R_s):")
    D.eqs(r"\Delta p_{ideal} = 2 f \frac{G_s^{2}}{\rho} N_{tcc} \left( \frac{\mu_b}{\mu_w} \right)^{-0.14}",
          r"\Delta p_{cross} = \Delta p_{ideal} (N_b - 1) R_b R_l",
          r"\Delta p_{window} = (2 + 0.6 N_{tcw}) \frac{G_w^{2}}{2\rho} N_b R_l, \qquad G_w = \frac{\dot{m}_c}{\sqrt{S_m S_w}}",
          r"\Delta p_{ends} = \Delta p_{ideal} \left( 1 + \frac{N_{tcw}}{N_{tcc}} \right) R_b R_s",
          r"\Delta p_{shell} = \Delta p_{cross} + \Delta p_{window} + \Delta p_{ends}", label="dpshell")

    D.h3("6.3.1  Two-phase pressure drop")

    D.p("Inside the dome the drop is referenced to the ALL-LIQUID Bell-Delaware value "
        "and scaled by the Chisholm two-phase multiplier, rather than evaluating "
        "Bell-Delaware at the homogeneous mixture density. The multiplier recovers both "
        "limits exactly \u2014 unity at x=0 and Gamma-squared at x=1:")
    D.eqs(r"\Gamma = \sqrt{\frac{\rho_l}{\rho_v}} \left( \frac{\mu_v}{\mu_l} \right)^{0.1}",
          r"\phi^{2} = 1 + (\Gamma^{2} - 1) \left[ B\, x^{(2-n)/2} (1-x)^{(2-n)/2} + x^{2-n} \right]",
          r"\Delta p_{TP} = \phi^{2}\, \Delta p_{liquid}",
          label="chisholm")
    D.p("with n = 0.2 for turbulent flow. B is selected on (Gamma, mass flux); at this "
        "exchanger's water design point Gamma = 3.60 and G ~ 3000 kg/m2s, giving B ~ 1.01.")

    D.note("B is reconstructed, not transcribed",
           "Chisholm (1973) and Grant & Chisholm (1979) were unavailable when this was "
           "implemented. The B table is corroborated at exactly one point \u2014 Hellborg "
           "(2017) eq. 137 hardcodes B = 21/Gamma, which lands precisely in the "
           "9.5 < Gamma < 28 branch. That is encouraging, not verification. Note also that "
           "Hellborg's hardcoded branch is WRONG for this exchanger: at Gamma = 3.60 the "
           "correct branch gives B ~ 1.01, not 5.83, so copying that equation verbatim "
           "would inflate the multiplier roughly sixfold.",
           colour=FLAG)

    D.h3("6.3.2  Momentum (acceleration) pressure drop")

    D.p("As the coolant vaporizes its density collapses and the flow must accelerate, "
        "which costs pressure over and above wall friction. Across a cell:")
    D.eq(r"\Delta p_{acc} = G^{2} \left[ \left( \frac{1}{\rho} \right)_{i+1} - \left( \frac{1}{\rho} \right)_{i} \right]",
         label="dpacc")
    D.p("This is the homogeneous form: the separated-flow momentum term "
        "x\u00b2/(alpha*rho_v) + (1-x)\u00b2/((1-alpha)*rho_l) reduces to 1/rho for the "
        "homogeneous void fraction the state closure returns, so a drift-flux alpha would "
        "be needed to go further. The term had been omitted entirely; it is worth "
        "**26 % of the water design point's total shell-side drop** (2.99 of 11.05 bar, "
        "where density falls 1000 \u2192 29 kg/m3) and 23 % of the N2 point's.")

    D.note("Correlation set is a matched pair",
           "The ideal-bank j and f coefficient tables and the J/R correction factors were calibrated "
           "against one another. Substituting a different tube-bank correlation without re-deriving "
           "the corrections invalidates the calibration, even though each component is individually "
           "defensible.")

    D.h2("6.4  Regime dispatch")

    D.p("Bell-Delaware contains no boiling physics whatsoever, so inside the two-phase dome it must "
        "be abandoned outright. At supercritical pressure the situation is more subtle: what the "
        "property-ratio closures correct for is the steep property variation near the "
        "pseudo-critical temperature T_pc(p) — the smeared-out remnant of the latent-heat spike. "
        "Away from that region a supercritical fluid is an ordinary single-phase fluid, and the "
        "cross-flow-calibrated Bell-Delaware correlation is the better-matched model for a baffled "
        "bundle. Selection is therefore made on the local state:")

    D.table(
        ["Local coolant regime", "Heat-transfer closure", "Friction closure"],
        [["Subcooled liquid or superheated vapour (single phase, subcritical)",
          "Bell-Delaware, " + D.ref("hshell"), "Bell-Delaware, " + D.ref("dpshell")],
         ["Saturated two-phase, 0 ≤ x ≤ 1",
          "Gungor-Winterton, " + D.ref("gw"), "Müller-Steinhagen & Heck, " + D.ref("msh")],
         ["Supercritical, bulk-to-wall interval reaching the pseudo-critical band around T_pc(p)",
          "Property-ratio closure from the registry (McCarthy-Wolf / Taylor)",
          "Registry friction gradient"],
         ["Supercritical, but far from T_pc — no property-ratio correction warranted",
          "Bell-Delaware, " + D.ref("hshell"),
          "Registry friction gradient"]],
        "Shell-side closure dispatch. Regime is evaluated per node, per sweep, from the running "
        "(p, h) state of the coolant and the lagged wall temperature. Note the last row: the "
        "heat-transfer closure falls back to Bell-Delaware while the friction model does not — "
        "see below.",
        widths=[6.4, 6.2, 3.8])

    D.note("Supercritical pressure alone is not the criterion",
           "Selection used to be gated purely on the coolantProp.coolant_model string, so which "
           "closure a fluid received was a configuration artefact rather than a statement about its "
           "state. Helium and N2 are both supercritical yet sit in entirely different places: "
           "Helium at 80 bar has T_pc = 11.4 K against a 300-1400 K march — 26 to 120x above any "
           "critical anomaly, with cp flat to 0.1 % — whereas N2 at 88 bar has T_pc = 147.8 K, a "
           "bulk of 100-124 K and a wall reaching 164 K, placing the pseudo-critical transition "
           "inside the thermal boundary layer. At the N2 design point 97 of 200 nodes meet the "
           "test. The selection is latched per node (one-way) so it cannot oscillate while the "
           "lagged wall temperature is still rising off its cold seed.",
           colour=FLAG)

    D.note("The pressure march stays on one friction model",
           "Bell-Delaware returns a LUMPED whole-bundle pressure drop, apportioned as dp_shell/N, "
           "whereas the registry closure returns a LOCAL gradient. Mixing the two across nodes "
           "double-counts, because each Bell-Delaware node contributes a share of a whole-bundle "
           "drop the other nodes are not taking. A supercritical node that falls back to "
           "Bell-Delaware for its heat-transfer coefficient therefore still takes its pressure drop "
           "from the local gradient. On this geometry Bell-Delaware's pressure drop is in any case "
           "not usable — see Section 11.")

    D.h2("6.5  Saturated flow boiling")

    D.p("Within the two-phase dome the Gungor-Winterton (1986) correlation is used. Its two-term "
        "structure carries the physics: forced convection is enhanced by vapour acceleration, while "
        "nucleate boiling is suppressed as that convection thins the near-wall superheated layer.")
    D.eqs(r"Re_l = \frac{G(1-x)D}{\mu_l}, \qquad h_l = 0.023\, Re_l^{0.8} Pr_l^{0.4} \frac{k_l}{D}",
          r"X_{tt} = \left( \frac{1-x}{x} \right)^{0.9} \left( \frac{\rho_v}{\rho_l} \right)^{0.5} \left( \frac{\mu_l}{\mu_v} \right)^{0.1}",
          r"Bo = \frac{q''}{G\, h_{fg}}",
          r"E = 1 + 24000\, Bo^{1.16} + 1.37 \left( \frac{1}{X_{tt}} \right)^{0.86}",
          r"S = \frac{1}{1 + 1.15 \cdot 10^{-6} E^{2} Re_l^{1.17}}",
          r"h_{pool} = 55\, p_r^{0.12} (-\log_{10} p_r)^{-0.55} M^{-0.5} q''^{0.67}",
          r"h = E\, h_l + S\, h_{pool}", label="gw")

    D.p("The pool-boiling term is the Cooper correlation, with M the molar mass in kg/kmol. The "
        "horizontal low-Froude correction of the original reference is not applied; the correlation "
        "is used here in its vertical / high-Froude form.")

    D.h2("6.6  Two-phase friction")

    D.p("The Müller-Steinhagen and Heck correlation interpolates empirically between the "
        "all-liquid and all-vapour limits, recovering each correctly at x = 0 and x = 1:")
    D.eqs(r"A = \frac{f_{l0} G^{2}}{2 D \rho_l}, \qquad B = \frac{f_{v0} G^{2}}{2 D \rho_v}",
          r"-\frac{dp}{dz} = \left[ A + 2(B-A)x \right] (1-x)^{1/3} + B x^{3}", label="msh")

    D.h2("6.7  Critical heat flux and onset of nucleate boiling")

    D.p("Departure from nucleate boiling is monitored through a lookup-table critical heat flux "
        "(Groeneveld 2006), reported as a margin:")
    D.eq(r"\text{CHF margin} = \frac{q''_{CHF}}{q''_w}")

    D.note("Undefined outside the dome",
           "The CHF margin is returned as NaN for subcooled liquid, superheated vapour and "
           "supercritical states. This is a genuine regime boundary — no critical-heat-flux "
           "concept exists there — and not missing data. A margin that falls to zero exactly at "
           "x → 1 denotes the dryout transition into complete vaporization, which is the intended "
           "behaviour for a design that deliberately superheats its coolant.")

    D.p("Two distinct mechanisms guard the subcooled-to-boiling transition and should not be "
        "conflated. The first is a **numerical blend window** of half-width 0.02 in quality, centred "
        "on x = 0, which smooths an otherwise abrupt change in heat-transfer coefficient between the "
        "single-phase and two-phase branches. It is a smoothing device with no physical content. The "
        "second is the **physical** criterion, the Bergles-Rohsenow (1964) onset-of-nucleate-boiling "
        "wall superheat, evaluated as a diagnostic and reported when exceeded:")
    D.eq(r"q''_{ONB} = 1082\, p^{1.156} \left( 1.8\, \Delta T_{ONB} \right)^{2.16/p^{0.0234}}")

    D.p("with p in bar, inverted for ΔT_ONB and compared against the estimated wall superheat. "
        "Subcooled nucleate boiling can begin well before the bulk fluid reaches saturation if the "
        "wall is sufficiently hot; when the solver reports an ONB exceedance it is stating that the "
        "smooth numerical blend is standing in for boiling the physics indicates has already begun.")

    # ---------------- 7 ----------------
    D.h1("7.  Shell-side energy march")

    D.p("The nodal duty computed per representative tube is scaled to the full bundle once, here:")
    D.eq(r"dQ_{total,i} = dQ_i\, N_{tubes}", label="ntubes")

    D.h2("7.1  Single-phase (temperature) march")
    D.eq(r"T_{i+1} = T_i + \frac{dQ_{total,i}}{\dot{m}_c\, c_{p,c}(T_i)}", label="Tmarch")

    D.h2("7.2  Real-fluid (enthalpy) march")
    D.eqs(r"h_{i+1} = h_i + \frac{dQ_{total,i}}{\dot{m}_c}",
          r"p_{i+1} = \max \left( p_i - \Delta p_i^{lagged},\; 1 \right)",
          r"(T, x, \alpha)_i = \text{flash}_{ph}(p_i, h_i)", label="hmarch")

    D.note("Why enthalpy rather than temperature",
           "A cp·ΔT update has nowhere to place latent heat. Marched in temperature, boiling "
           "coolant would continue to rise past 1000 K instead of pinning near saturation while it "
           "vaporizes. The deeper distinction is which variable determines the thermodynamic state: "
           "for single-phase flow at fixed pressure, h = h(T) is monotonic and invertible, so the two "
           "marches carry identical information. Inside the dome, T = T_sat(p) is **degenerate** — "
           "identical at x = 0.1 and x = 0.9 — so temperature cannot represent the state at all, "
           "while (p, h) still can.",
           colour=FLAG)

    D.p("The single-phase march is a forward-Euler approximation to the enthalpy march, with cp "
        "evaluated at each node's starting temperature. It is exact for constant cp and first-order "
        "otherwise, with the error scaling as 1/N. For Helium, which is monatomic and has cp flat to "
        "roughly 0.1 % between 300 K and 1400 K at 80 bar, the resulting outlet-temperature error is "
        "immeasurably small. Section 10 quantifies this per fluid.")

    D.h2("7.3  Pressure march and lagging")

    D.p("The nodal pressure drop used in " + D.ref("hmarch") + " is taken from the **previous** sweep. This is "
        "unavoidable: the pressure profile depends on the enthalpy profile through local density and "
        "quality, while the enthalpy profile depends on the pressure profile through the local "
        "saturation temperature. Lagging one of the two breaks the circular dependency at first-order "
        "accuracy in the sweep iteration, and the lag vanishes at convergence.")

    D.h2("7.4  Equation-of-state ceiling handling")

    D.p("A very low coolant flow against a large duty can superheat the coolant past the property "
        "library's validity ceiling. This is a genuine physical edge case rather than a numerical "
        "failure. When the flash fails, the last valid state, enthalpy and pressure triple is held "
        "for the remainder of the march. All three must be frozen together; freezing only the state "
        "while enthalpy and pressure continued to drift would re-trigger the same failure on the "
        "following sweep, when those drifted values are read back.")

    # ---------------- 8 ----------------
    D.h1("8.  Outer iteration (predictive sweep)")

    D.figure(os.path.join(fig_dir, "fig_sweep.png"), 16.2,
             "The predictive sweep. The two marches exchange a duty field forward and a temperature "
             "field back. Three closure inputs cannot be evaluated before the march that produces them "
             "and are therefore consumed one sweep behind.")

    D.p("The two marches are closed by an under-relaxed Picard (fixed-point) iteration:")
    D.eqs(r"tube \leftarrow \text{march}(T_{shell}^{(k)})",
          r"T_{new} \leftarrow \text{shell march}(dQ)",
          r"\delta = \max \left| T_{new} - T_{shell}^{(k)} \right|",
          r"T_{shell}^{(k+1)} = (1-\omega) T_{shell}^{(k)} + \omega\, T_{new}")

    D.p("Iteration continues until δ < 0.05 K, to a maximum of 25 sweeps. Under-relaxation at "
        "ω = 0.5 is functionally necessary rather than cosmetic: the coupling is stiff, since the "
        "coolant film coefficient depends on the shell temperature, which depends on the duty, which "
        "in turn depends on the film coefficient. An unrelaxed update oscillates or diverges, "
        "particularly across the boiling transition where the sensitivity of the film coefficient to "
        "shell temperature is greatest.")

    D.h2("8.1  Lagged closure quantities")

    D.table(
        ["Quantity", "Consumed by", "Seeded at", "Reason for lag"],
        [["q″_w", "Boiling-number term in Gungor-Winterton", "0",
          "Wall flux is an output of the tube-side march that the shell-side closure needs as input"],
         ["Δp per node", "Shell-side pressure march", "0",
          "Pressure and enthalpy profiles are mutually dependent"],
         ["T_wc", "Property ratio in supercritical closures", "coolant inlet T",
          "Coolant-side wall temperature is an output of the wall solve"]],
        "Quantities lagged by one sweep iteration. All are seeded neutrally and all converge to "
        "self-consistency as δ → 0.",
        widths=[2.6, 5.4, 2.6, 5.8], mono_cols=(0,))

    D.h2("8.2  Initialization")

    D.note("Seeding differs by coolant model",
           "Single-phase mode can seed the shell field with a linear profile between the two "
           "prescribed inlets. Real-fluid mode cannot: a prescribed coolant outlet temperature is "
           "meaningless when the coolant may be two-phase, because inside the dome T and p are not "
           "independent and a single (T, p) pair cannot represent the state. Real-fluid mode therefore "
           "seeds uniformly at the physical subcooled inlet condition.")

    D.h2("8.3  Convergence behaviour")

    D.p("Convergence is monotone for well-conditioned single-phase cases. Near a boiling transition "
        "the iteration may instead settle into a limit cycle of order 0.5–1 K amplitude without "
        "driving δ below tolerance. The reported outlet is stable to within that band, but this is "
        "a genuine non-convergence and is reported as such rather than being silently accepted. Where "
        "a design point exhibits this behaviour it should be recorded alongside the result.")

    D.h2("8.4  Global energy closure")

    D.eq(r"Q_{tot} = \sum_{i}{ dQ_i\, N_{tubes} }")

    D.p("The gas side is always enthalpy-marched, so the duty subtracted from the gas budget is "
        "exactly what crosses the wall. Whether the coolant side closes exactly depends on which "
        "march of Section 7 was used:")
    D.bullet("**Real-fluid (enthalpy) mode — exact by construction.** Enthalpy is the marched "
             "variable, so the coolant absorbs precisely the summed duty by definition. No closure "
             "error exists to monitor.")
    D.bullet("**Single-phase (temperature) mode — first order.** Temperature is accumulated using "
             "cp at each node's start, so the true enthalpy rise and the summed duty differ in "
             "general. The discrepancy is zero for constant cp and diminishes as 1/N.")
    D.p("In neither case is there time-integration drift to accumulate, which is why the steady "
        "solver carries no energy-residual diagnostic, unlike the transient solvers where such a "
        "diagnostic is essential.")

    # ---------------- 9 ----------------
    D.h1("9.  Structural check")

    D.p("The tube loading in this configuration is the reverse of a regeneratively cooled channel: "
        "high-pressure coolant surrounds the tubes while comparatively low-pressure gas flows inside. "
        "The tubes are therefore under **external** pressure and the failure mode of interest is "
        "elastic collapse rather than burst. This check is invoked separately from the thermal solve.")
    D.eqs(r"\Delta P_{ext} = p_{c,in} - p_0",
          r"\sigma_{ext} = \sigma_{ext}(\Delta P_{ext}, s_w, D_i)",
          r"P_{cr} = P_{cr}(E(\bar{T}), s_w, D_i, \nu)",
          r"\text{collapse margin} = \frac{\left| \Delta P_{ext} \right|}{P_{cr}} \; < 1 \;\text{ for safety}")

    D.p("Young's modulus is evaluated at the mean of the maximum gas-side and minimum coolant-side "
        "wall temperatures, converted to degrees Celsius, since the material property tables are "
        "tabulated in °C while the solver is otherwise SI throughout.")

    D.note("Extrapolation beyond the characterized range",
           "Inconel 718 property data spans −240 °C to 760 °C (1033.15 K). Above that "
           "limit the yield curve clamps flat rather than extrapolating a physical material law, so a "
           "stress ratio computed at, say, 1100 K is arithmetically well-formed but not physically "
           "supported. Maximum wall temperature must always be reported alongside the collapse margin "
           "so that this condition is visible.",
           colour=FLAG)

    # ---------------- 10 ----------------
    D.h1("10.  Application to Helium, LN₂/N₂ and Water")

    D.p("The three coolants of interest exercise different branches of the model. The gas-side, wall "
        "and iteration structure is identical in all three cases; only the shell-side closure and the "
        "energy-march variable differ. This section is the bridge between the framework above and the "
        "reported results.")

    D.table(
        ["Coolant", "Typical regime at design point", "Shell-side closure", "Energy march", "Status"],
        [["Helium",
          "Single phase throughout, well above critical pressure but gas-like in behaviour",
          "Bell-Delaware",
          D.ref("Tmarch"),
          "**Baseline.** cp is flat to ≈ 0.1 % over the duty range, so the temperature march is "
          "numerically indistinguishable from an enthalpy march."],
         ["LN₂ / supercritical N₂",
          "Supercritical throughout at any pressure able to meet a 75 bar exit target, since p_crit "
          "≈ 34 bar. Quality undefined; liquid-like dense fluid.",
          "Registry property-ratio closure (McCarthy-Wolf / Taylor)",
          D.ref("hmarch"),
          "**Bounded sensitivity.** The supercritical correlations remain straight-tube fits applied "
          "to cross-flow, but substituting Bell-Delaware moves Q_tot by only 1 % (the coolant film "
          "is 0.58 % of total resistance); wall temperature carries ~54 K of uncertainty."],
         ["Water",
          "Traverses subcooled liquid, the two-phase dome, and superheated vapour within a single "
          "exchanger when the design targets a steam outlet.",
          "Bell-Delaware → Gungor-Winterton → Bell-Delaware, with blending at the dome entry",
          D.ref("hmarch"),
          "**Partly validated.** The boiling correlations are literature-validated for tube flow; "
          "cross-flow enhancement of boiling is real and not modelled."]],
        "Coolant-specific model paths. The common caveat for the two real-fluid cases is that "
        "Gungor-Winterton, Müller-Steinhagen-Heck and the supercritical property-ratio closures are "
        "all tube-flow correlations applied to shell-side cross-flow.",
        widths=[2.2, 4.4, 3.4, 2.2, 4.2])

    D.note("Cross-flow boiling enhancement",
           "Shell-side cross-flow over a tube bundle enhances boiling heat transfer relative to flow "
           "inside a tube, through mechanisms none of the correlations used here represent. The model "
           "is therefore expected to be **conservative** on coolant-side heat-transfer coefficient in "
           "the two-phase and supercritical branches. This is a known open item, not a modelling "
           "oversight, and it is surfaced by the solver at runtime rather than left silent.",
           colour=FLAG)

    # ---------------- 11 ----------------
    D.h1("11.  Assumptions and validity")

    D.table(
        ["Assumption", "Rationale", "Where it becomes significant"],
        [["Bell-Delaware baffle-leakage ratio r_lm is far outside the correlation's fitted range",
          "r_lm = (S_sb+S_tb)/S_m = 6.5 on this geometry against a fitted range of roughly r_lm <= 1: "
          "the tube-to-baffle clearance area (17.7 cm2) is 6.5x the cross-flow area (2.9 cm2)",
          "**Both leakage corrections are extrapolated**: J_l = 0.41 on heat transfer, and R_l = 0.009 "
          "on pressure drop. R_l crushes the cross-flow and window terms, leaving the uncorrected "
          "end-zone term to dominate a whole-bundle drop of ~288 bar. Shell-side pressure drop from "
          "Bell-Delaware is **not usable** on this geometry; reported exit pressures come from the "
          "local friction-gradient model. Affects every Bell-Delaware result including the Helium "
          "baseline. Reported at runtime."],
         ["Two-phase and supercritical closures are tube-flow correlations applied to cross-flow",
          "No validated shell-side cross-flow boiling or supercritical closure is available",
          "Any two-phase shell-side case. For the supercritical N2 design point the sensitivity has "
          "been measured: substituting Bell-Delaware changes h_c by 3x but Q_tot by only 1 %, because "
          "the coolant film carries just 0.58 % of the total thermal resistance (gas film 95.2 %, "
          "wall 4.2 %). Wall temperature carries ~54 K of closure uncertainty."],
         ["Shell-side pressure drop apportioned uniformly across axial nodes (single-phase branch)",
          "Bell-Delaware returns a lumped whole-bundle drop; a per-baffle split would require its "
          "per-crossing terms to be re-derived",
          "When local saturation temperature matters and the true drop is strongly non-uniform axially"],
         ["Closure quantities lagged by one sweep",
          "Breaks an otherwise circular dependency between the two marches",
          "First order in the sweep; vanishes at convergence but contributes to limit cycling where "
          "convergence stalls"],
         ["Gas pressure does not feed back into gas density",
          "The FPV manifold is constructed at a fixed reference pressure",
          "Only if gas Mach number rises materially above the 0.3 design ceiling"],
         ["Shell-side single-phase branch holds pressure at the inlet value",
          "Enthalpy of a near-ideal gas is weakly pressure-dependent",
          "Density and transport properties do vary with pressure; the computed Bell-Delaware drop is "
          "presently discarded in this branch"],
         ["Shell side is zero-dimensional per node with no momentum ODE",
          "Baffled cross-flow has no single streamwise coordinate; this is the appropriate model form",
          "Not a limitation in the modelled regime"],
         ["Gas radiation disabled",
          "Model implemented but switched off in this solver path",
          "Very high gas temperature combined with high gas emissivity"],
         ["Ideal-gas relations in the legacy quasi-1D coolant equations",
          "Applies only to the helical-coil solver, not to this one",
          "Not applicable to shell-and-tube"]],
        "Model assumptions and their domains of validity, ordered approximately by practical impact "
        "on the results reported in this programme.",
        widths=[5.0, 5.6, 5.8])

    D.h2("11.1  Self-reporting diagnostics")

    D.p("Several of the limitations above surface at runtime rather than remaining implicit in "
        "documentation. When reading a solver run, the following messages are load-bearing and should "
        "be recorded with the result:")
    D.bullet("Supercritical cross-flow geometry extrapolation warning, naming the selected closure.")
    D.bullet("Bergles-Rohsenow onset-of-nucleate-boiling exceedance, with the wall-superheat margin.")
    D.bullet("Sweep non-convergence notice, with the final residual.")
    D.bullet("NaN rather than a fabricated value for CHF margin outside the two-phase dome.")
    D.bullet("Regimes traversed and count of heat-transfer-deterioration risk nodes in supercritical operation.")
    D.bullet("Bell-Delaware baffle-leakage ratio r_lm exceeding the correlation's fitted range.")
    D.bullet("Count of supercritical nodes whose heat-transfer coefficient came from Bell-Delaware "
             "rather than a property-ratio closure.")

    return D


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    fig_dir = os.environ.get("SHELLTUBE_FIG_DIR", here)
    out = os.path.join(here, "SHELLTUBE_STEADY_THEORY.docx")
    first = build(fig_dir)               # pass 1: collect equation labels
    doc = build(fig_dir, first.labels)   # pass 2: resolve cross-references
    doc.save(out)
    missing = [k for k, v in first.labels.items() if v is None]
    if missing:
        print("  ! unresolved equation labels: " + ", ".join(missing), file=sys.stderr)
    print(f"wrote {out}")
    print(f"  equations: {doc.eq_n}   figures: {doc.fig_n}   tables: {doc.tab_n}")


if __name__ == "__main__":
    main()
